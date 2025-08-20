#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF转换引擎API
支持PDF与Word、图片等格式的相互转换
"""

import os
import json
import logging
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings
from django.contrib.auth.decorators import login_required
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False
    print("警告: PyMuPDF (fitz) 未安装，PDF转换功能将受限")
from PIL import Image
import io
import base64
from datetime import datetime
import uuid

# 配置日志
logger = logging.getLogger(__name__)

class PDFConverter:
    """PDF转换引擎核心类"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': ['.pdf'],
            'word': ['.doc', '.docx'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff'],
            'text': ['.txt']
        }
    
    def validate_file(self, file, expected_type):
        """验证文件格式"""
        if not file:
            return False, "文件不能为空"
        
        # 检查file对象是否有name属性
        if not hasattr(file, 'name') or not file.name:
            return False, "无效的文件对象"
        
        file_ext = os.path.splitext(file.name)[1].lower()
        
        # 检查文件大小 (限制为50MB)
        if hasattr(file, 'size') and file.size > 50 * 1024 * 1024:
            return False, "文件大小不能超过50MB"
        
        # 检查文件格式兼容性
        if expected_type in self.supported_formats:
            if file_ext not in self.supported_formats[expected_type]:
                # 提供智能提示和自动切换建议
                suggestion = self._get_conversion_suggestion(file_ext, expected_type)
                return False, suggestion
        
        return True, "文件验证通过"
    
    def _get_conversion_suggestion(self, file_ext, current_type):
        """获取转换类型建议"""
        # 检查file_ext是否为字符串
        if not isinstance(file_ext, str):
            return f"无效的文件扩展名: {file_ext}"
        
        # 定义文件类型到转换类型的映射
        file_type_mapping = {
            '.pdf': 'pdf',
            '.doc': 'word',
            '.docx': 'word',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.png': 'image',
            '.bmp': 'image',
            '.tiff': 'image',
            '.gif': 'image',
            '.txt': 'text'
        }
        
        # 定义转换类型到操作类型的映射
        conversion_mapping = {
            'pdf': {
                'pdf-to-word': 'PDF转Word',
                'pdf-to-image': 'PDF转图片'
            },
            'word': {
                'word-to-pdf': 'Word转PDF'
            },
            'image': {
                'image-to-pdf': '图片转PDF'
            },
            'text': {
                'text-to-pdf': '文本转PDF'
            }
        }
        
        # 获取文件的实际类型
        actual_type = file_type_mapping.get(file_ext, 'unknown')
        
        if actual_type == 'unknown':
            return f"不支持的文件格式: {file_ext}。请使用支持的文件格式。"
        
        # 获取当前转换类型的显示名称
        current_display = conversion_mapping.get(actual_type, {}).get(current_type, current_type)
        
        # 获取建议的转换类型
        suggested_conversions = conversion_mapping.get(actual_type, {})
        
        if not suggested_conversions:
            return f"文件格式 {file_ext} 不支持任何转换操作。"
        
        # 构建建议信息
        suggestion = f"文件格式 {file_ext} 与当前转换类型 '{current_display}' 不兼容。\n\n"
        suggestion += "建议的转换类型：\n"
        
        for conv_type, display_name in suggested_conversions.items():
            suggestion += f"• {display_name} ({conv_type})\n"
        
        suggestion += f"\n请切换到适合的转换类型，或上传 {actual_type} 格式的文件。"
        
        return suggestion
    
    def pdf_to_word(self, pdf_file):
        """PDF转Word - 真实实现"""
        try:
            # 检查pdf2docx库是否可用
            try:
                from pdf2docx import Converter
            except ImportError:
                return False, "pdf2docx库未安装，无法进行PDF转Word转换", None
            
            # 重置文件指针
            pdf_file.seek(0)
            
            # 使用pdf2docx进行真实转换
            import tempfile
            import os
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                temp_pdf.write(pdf_file.read())
                temp_pdf_path = temp_pdf.name
            
            # 创建临时输出文件路径
            temp_docx_path = temp_pdf_path.replace('.pdf', '.docx')
            
            try:
                # 检查输入PDF文件是否存在
                if not os.path.exists(temp_pdf_path):
                    return False, "临时PDF文件创建失败", None
                
                # 使用pdf2docx进行转换
                cv = Converter(temp_pdf_path)
                cv.convert(temp_docx_path)
                cv.close()
                
                # 检查输出文件是否存在
                if not os.path.exists(temp_docx_path):
                    return False, "转换失败：输出Word文件未生成", None
                
                # 读取转换后的文件
                with open(temp_docx_path, 'rb') as docx_file:
                    docx_content = docx_file.read()
                
                # 清理临时文件
                try:
                    os.unlink(temp_pdf_path)
                    os.unlink(temp_docx_path)
                except:
                    pass
                
                if len(docx_content) == 0:
                    return False, "转换后的文件为空，可能是扫描版PDF或内容无法识别", None
                
                # 检查转换结果是否包含实际内容
                try:
                    from docx import Document
                    doc = Document(io.BytesIO(docx_content))
                    text_content = ""
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                    
                    if len(text_content.strip()) < 10:  # 如果文本内容太少，可能是扫描版PDF
                        return False, "检测到扫描版PDF，无法提取文本内容。请使用OCR工具处理。", None
                        
                except Exception as check_error:
                    logger.warning(f"转换结果检查失败: {check_error}")
                    # 继续处理，不因为检查失败而中断
                
                return True, docx_content, "pdf_to_word"
                
            except Exception as conversion_error:
                # 清理临时文件
                try:
                    if os.path.exists(temp_pdf_path):
                        os.unlink(temp_pdf_path)
                    if os.path.exists(temp_docx_path):
                        os.unlink(temp_docx_path)
                except:
                    pass
                
                # 提供更详细的错误信息
                error_msg = str(conversion_error)
                if "Package not found" in error_msg:
                    return False, "PDF文件损坏或格式不支持，请检查文件完整性", None
                elif "Permission denied" in error_msg:
                    return False, "文件访问权限不足，请检查文件权限", None
                else:
                    return False, f"PDF转Word转换失败: {error_msg}", None
            
        except Exception as e:
            logger.error(f"PDF转Word失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None
    
    def word_to_pdf(self, word_file):
        """Word转PDF - 真实实现"""
        try:
            # 重置文件指针
            word_file.seek(0)
            
            # 使用python-docx和reportlab进行转换
            import tempfile
            import os
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                temp_docx.write(word_file.read())
                temp_docx_path = temp_docx.name
            
            # 创建临时输出文件路径
            temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
            
            try:
                # 使用python-docx和reportlab进行转换
                from docx import Document
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                
                # 检查临时文件是否存在
                if not os.path.exists(temp_docx_path):
                    return False, "临时Word文件创建失败", None
                
                # 读取Word文档
                doc = Document(temp_docx_path)
                
                # 创建PDF
                c = canvas.Canvas(temp_pdf_path, pagesize=letter)
                y = 750  # 起始Y坐标
                
                # 检查文档是否有内容
                if not doc.paragraphs:
                    c.drawString(72, y, "空文档")
                else:
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            # 处理长文本换行
                            text = paragraph.text
                            words = text.split()
                            line = ""
                            
                            for word in words:
                                test_line = line + " " + word if line else word
                                if len(test_line) * 7 < 500:  # 简单的字符宽度估算
                                    line = test_line
                                else:
                                    if line:
                                        c.drawString(72, y, line)
                                        y -= 20
                                        line = word
                                        
                                        if y < 50:  # 如果页面空间不足，添加新页面
                                            c.showPage()
                                            y = 750
                            
                            if line:
                                c.drawString(72, y, line)
                                y -= 20
                                
                                if y < 50:  # 如果页面空间不足，添加新页面
                                    c.showPage()
                                    y = 750
                
                c.save()
                
                # 检查输出文件是否存在
                if not os.path.exists(temp_pdf_path):
                    return False, "转换失败：输出PDF文件未生成", None
                
                # 读取转换后的文件
                with open(temp_pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # 清理临时文件
                try:
                    os.unlink(temp_docx_path)
                    os.unlink(temp_pdf_path)
                except:
                    pass
                
                if len(pdf_content) == 0:
                    return False, "转换后的文件为空，可能是Word文档内容无法识别", None
                
                return True, pdf_content, "word_to_pdf"
                
            except Exception as conversion_error:
                # 清理临时文件
                try:
                    if os.path.exists(temp_docx_path):
                        os.unlink(temp_docx_path)
                    if os.path.exists(temp_pdf_path):
                        os.unlink(temp_pdf_path)
                except:
                    pass
                
                # 提供更详细的错误信息
                error_msg = str(conversion_error)
                if "Package not found" in error_msg:
                    return False, "Word文件损坏或格式不支持，请检查文件完整性", None
                elif "Permission denied" in error_msg:
                    return False, "文件访问权限不足，请检查文件权限", None
                else:
                    return False, f"Word转PDF转换失败: {error_msg}", None
            
        except Exception as e:
            logger.error(f"Word转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None
    
    def pdf_to_images(self, pdf_file, dpi=150):
        """PDF转图片"""
        try:
            if not FITZ_AVAILABLE:
                return False, "PyMuPDF未安装，无法进行PDF转换", None
            
            doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
            images = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                mat = fitz.Matrix(dpi/72, dpi/72)  # 设置DPI
                pix = page.get_pixmap(matrix=mat)
                
                # 转换为PIL Image
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                # 转换为base64
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
                
                images.append({
                    'page': page_num + 1,
                    'data': img_base64,
                    'width': img.width,
                    'height': img.height
                })
            
            doc.close()
            return True, images, "pdf_to_images"
            
        except Exception as e:
            logger.error(f"PDF转图片失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None
    
    def images_to_pdf(self, image_files):
        """图片转PDF"""
        try:
            images = []
            
            for img_file in image_files:
                img = Image.open(img_file)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                images.append(img)
            
            if not images:
                return False, "没有有效的图片文件", None
            
            # 创建PDF
            pdf_buffer = io.BytesIO()
            if len(images) == 1:
                images[0].save(pdf_buffer, format='PDF')
            else:
                images[0].save(pdf_buffer, format='PDF', save_all=True, append_images=images[1:])
            
            pdf_content = pdf_buffer.getvalue()
            return True, pdf_content, "images_to_pdf"
            
        except Exception as e:
            logger.error(f"图片转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None
    
    def text_to_pdf(self, text_content):
        """文本转PDF"""
        try:
            # 检查reportlab库是否可用
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.units import inch
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            except ImportError:
                return False, "reportlab库未安装，无法进行文本转PDF转换", None
            
            # 创建PDF缓冲区
            pdf_buffer = io.BytesIO()
            
            # 创建PDF文档
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            story = []
            
            # 获取样式
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            
            # 设置中文字体支持
            try:
                # 使用reportlab内置的中文字体
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                normal_style.fontName = 'STSong-Light'
                normal_style.fontSize = 12
                normal_style.leading = 14
                normal_style.alignment = 0  # 左对齐
            except Exception as e:
                try:
                    # 尝试使用系统中文字体
                    import platform
                    if platform.system() == 'Darwin':  # macOS
                        try:
                            pdfmetrics.registerFont(TTFont('PingFang', '/System/Library/Fonts/PingFang.ttc'))
                            normal_style.fontName = 'PingFang'
                        except:
                            pdfmetrics.registerFont(TTFont('HiraginoSans', '/System/Library/Fonts/STHeiti Light.ttc'))
                            normal_style.fontName = 'HiraginoSans'
                    elif platform.system() == 'Windows':
                        pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
                        normal_style.fontName = 'SimSun'
                    else:  # Linux
                        pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
                        normal_style.fontName = 'DejaVuSans'
                    
                    normal_style.fontSize = 12
                    normal_style.leading = 14
                    normal_style.alignment = 0
                except Exception as e2:
                    # 如果都不可用，使用默认字体
                    normal_style.fontName = 'Helvetica'
                    normal_style.fontSize = 12
                    normal_style.leading = 14
                    normal_style.alignment = 0
                    logger.warning(f"无法加载中文字体，将使用默认字体: {str(e2)}")
            
            # 处理文本内容
            lines = text_content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    # 创建段落
                    paragraph = Paragraph(line, normal_style)
                    story.append(paragraph)
                    story.append(Spacer(1, 6))  # 添加间距
                else:
                    # 空行
                    story.append(Spacer(1, 12))
            
            # 生成PDF
            doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            if len(pdf_content) == 0:
                return False, "生成的PDF文件为空", None
            
            return True, pdf_content, "text_to_pdf"
            
        except Exception as e:
            logger.error(f"文本转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def pdf_to_text(self, pdf_file):
        """PDF转文本"""
        try:
            if not FITZ_AVAILABLE:
                return False, "PyMuPDF (fitz) 库未安装，无法进行PDF转文本转换", None
            
            # 创建临时文件
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
                # 写入PDF内容
                for chunk in pdf_file.chunks():
                    temp_pdf.write(chunk)
                temp_pdf_path = temp_pdf.name
            
            try:
                # 打开PDF文件
                doc = fitz.open(temp_pdf_path)
                
                # 检查PDF是否为空或损坏
                if len(doc) == 0:
                    doc.close()
                    return False, "PDF文件为空或损坏", None
                
                text_content = ""
                total_pages = len(doc)
                
                # 逐页提取文本
                for page_num in range(total_pages):
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        text_content += page_text
                        if page_num < total_pages - 1:
                            text_content += "\n\n"  # 页面间添加空行
                    except Exception as page_error:
                        logger.warning(f"提取第{page_num + 1}页文本时出错: {str(page_error)}")
                        text_content += f"\n[第{page_num + 1}页文本提取失败]\n"
                
                doc.close()
                
                # 检查提取的文本内容
                if not text_content.strip():
                    return False, "PDF文件不包含可提取的文本内容（可能是扫描版PDF，建议使用OCR工具）", None
                
                # 如果文本内容很少，可能是扫描版PDF
                if len(text_content.strip()) < 10:
                    return False, "提取的文本内容过少，可能是扫描版PDF，建议使用OCR工具", None
                
                return True, text_content, "pdf_to_text"
                
            finally:
                # 清理临时文件
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)
                    
        except Exception as e:
            logger.error(f"PDF转文本失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

    def txt_to_pdf(self, txt_file):
        """TXT文件转PDF"""
        try:
            # 读取txt文件内容
            txt_content = txt_file.read().decode('utf-8')
            
            # 调用text_to_pdf方法
            return self.text_to_pdf(txt_content)
            
        except Exception as e:
            logger.error(f"TXT文件转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None

@csrf_exempt
@require_http_methods(["POST"])
def pdf_converter_api(request):
    """PDF转换API主入口"""
    try:
        # 导入模型
        from .models.legacy_models import PDFConversionRecord
        import time
        
        # 创建转换器实例
        converter = PDFConverter()
        
        # 添加调试信息
        logger.info(f"PDF转换API请求: POST数据={dict(request.POST)}, FILES={list(request.FILES.keys())}")
        
        # 添加更详细的调试信息
        if 'file' in request.FILES:
            file = request.FILES['file']
            logger.info(f"上传文件信息: 名称={file.name}, 大小={file.size}, 类型={file.content_type}")
        
        conversion_type = request.POST.get('type', '')
        
        # 检查是否有文件上传（文本转PDF除外）
        if conversion_type != 'text-to-pdf':
            if 'file' not in request.FILES:
                logger.warning("PDF转换API: 没有上传文件")
                return JsonResponse({
                    'success': False,
                    'error': '没有上传文件'
                }, status=400)
            file = request.FILES['file']
        else:
            # 文本转PDF不需要文件上传
            file = None
        
        # 创建转换记录（如果用户已登录）
        conversion_record = None
        if request.user.is_authenticated:
            if conversion_type == 'text-to-pdf':
                # 文本转PDF的特殊处理
                text_content = request.POST.get('text_content', '')
                original_filename = f"text_content_{len(text_content)}_chars.txt"
                file_size = len(text_content.encode('utf-8'))
            else:
                original_filename = file.name
                file_size = file.size
            
            conversion_record = PDFConversionRecord.objects.create(
                user=request.user,
                conversion_type=conversion_type.replace('-', '_'),
                original_filename=original_filename,
                file_size=file_size,
                status='processing'
            )
        
        start_time = time.time()
        
        # 验证转换类型
        valid_types = ['pdf-to-word', 'word-to-pdf', 'pdf-to-image', 'image-to-pdf', 'text-to-pdf', 'pdf-to-text', 'txt-to-pdf']
        if conversion_type not in valid_types:
            return JsonResponse({
                'success': False,
                'error': f'不支持的转换类型: {conversion_type}'
            }, status=400)
        
        # 根据转换类型验证文件格式
        if conversion_type == 'text-to-pdf':
            # 文本转PDF不需要文件验证
            text_content = request.POST.get('text_content', '')
            if not text_content.strip():
                is_valid, message = False, "请输入要转换的文本内容"
            else:
                is_valid, message = True, "文本内容验证通过"
        elif conversion_type == 'pdf-to-word':
            is_valid, message = converter.validate_file(file, 'pdf')
        elif conversion_type == 'word-to-pdf':
            is_valid, message = converter.validate_file(file, 'word')
        elif conversion_type == 'pdf-to-image':
            is_valid, message = converter.validate_file(file, 'pdf')
        elif conversion_type == 'image-to-pdf':
            is_valid, message = converter.validate_file(file, 'image')
        elif conversion_type == 'pdf-to-text':
            is_valid, message = converter.validate_file(file, 'pdf')
        elif conversion_type == 'txt-to-pdf':
            is_valid, message = converter.validate_file(file, 'text')
        else:
            is_valid, message = True, "文件验证通过"
        
        if not is_valid:
            # 检查是否包含转换建议
            if "建议的转换类型" in message:
                # 解析建议的转换类型
                suggested_types = []
                lines = message.split('\n')
                for line in lines:
                    if line.strip().startswith('•'):
                        # 提取转换类型
                        conv_type = line.split('(')[1].split(')')[0] if '(' in line else None
                        if conv_type:
                            suggested_types.append(conv_type)
                
                return JsonResponse({
                    'success': False,
                    'error': message,
                    'suggested_types': suggested_types,
                    'needs_type_switch': True
                }, status=400)
            else:
                return JsonResponse({
                    'success': False,
                    'error': message
                }, status=400)
        
        # 执行转换
        if conversion_type == 'text-to-pdf':
            text_content = request.POST.get('text_content', '')
            success, result, file_type = converter.text_to_pdf(text_content)
        elif conversion_type == 'pdf-to-word':
            success, result, file_type = converter.pdf_to_word(file)
        elif conversion_type == 'word-to-pdf':
            success, result, file_type = converter.word_to_pdf(file)
        elif conversion_type == 'pdf-to-image':
            success, result, file_type = converter.pdf_to_images(file)
        elif conversion_type == 'image-to-pdf':
            success, result, file_type = converter.images_to_pdf([file])
        elif conversion_type == 'pdf-to-text':
            success, result, file_type = converter.pdf_to_text(file)
        elif conversion_type == 'txt-to-pdf':
            success, result, file_type = converter.txt_to_pdf(file)
        else:
            return JsonResponse({
                'success': False,
                'error': '未知的转换类型'
            }, status=400)
        
        # 计算转换时间
        conversion_time = time.time() - start_time
        
        if not success:
            # 更新转换记录为失败状态（如果存在）
            if conversion_record:
                conversion_record.status = 'failed'
                conversion_record.error_message = result
                conversion_record.conversion_time = conversion_time
                conversion_record.save()
            
            return JsonResponse({
                'success': False,
                'error': result
            }, status=500)
        
        # 保存转换结果
        output_filename = f"{uuid.uuid4()}_{conversion_type.replace('-', '_')}"
        
        if file_type == 'pdf_to_word':
            output_filename += '.docx'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        elif file_type == 'word_to_pdf':
            output_filename += '.pdf'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        elif file_type == 'pdf_to_images':
            # 创建ZIP文件包含所有图片
            import zipfile
            from io import BytesIO
            import base64
            
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for i, img_data in enumerate(result):
                    # 解码base64图片数据
                    img_bytes = base64.b64decode(img_data['data'])
                    # 添加到ZIP文件
                    zip_file.writestr(f'page_{i+1}.png', img_bytes)
            
            zip_content = zip_buffer.getvalue()
            zip_buffer.close()
            
            # 保存ZIP文件
            output_filename += '_images.zip'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(zip_content))
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
            
            # 更新转换记录为成功状态（如果存在）
            if conversion_record:
                conversion_record.status = 'success'
                conversion_record.output_filename = output_filename
                conversion_record.conversion_time = conversion_time
                conversion_record.download_url = download_url
                conversion_record.save()
            
            # 返回下载链接
            return JsonResponse({
                'success': True,
                'type': 'file',
                'download_url': download_url,
                'filename': output_filename,
                'original_filename': file.name,
                'file_size': len(zip_content),
                'total_pages': len(result),
                'message': f'已转换{len(result)}页，打包为ZIP文件供下载',
                'conversion_type': conversion_type
            })
        elif file_type == 'images_to_pdf':
            output_filename += '.pdf'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        elif file_type == 'text_to_pdf':
            output_filename += '.pdf'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        elif file_type == 'pdf_to_text':
            output_filename += '.txt'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result.encode('utf-8')))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        else:
            return JsonResponse({
                'success': False,
                'error': '未知的文件类型'
            }, status=500)
        
        # 更新转换记录为成功状态（如果存在）
        if conversion_record:
            conversion_record.status = 'success'
            conversion_record.output_filename = output_filename
            conversion_record.conversion_time = conversion_time
            conversion_record.download_url = download_url
            conversion_record.save()
        
        # 确定原始文件名
        if conversion_type == 'text-to-pdf':
            original_filename = '文本内容'
        elif conversion_type == 'pdf-to-text':
            original_filename = file.name
        else:
            original_filename = file.name
        
        return JsonResponse({
            'success': True,
            'type': 'file',
            'download_url': download_url,
            'filename': output_filename,
            'original_filename': original_filename,
            'conversion_type': conversion_type
        })
        
    except Exception as e:
        logger.error(f"PDF转换API错误: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f'服务器错误: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["GET"])
def pdf_converter_status(request):
    """获取转换状态和功能支持情况"""
    try:
        # 检查依赖库状态
        import sys
        from datetime import datetime
        
        # 检查各种库的可用性
        pdf2docx_available = False
        docx2pdf_available = False
        pil_available = False
        
        try:
            from pdf2docx import Converter
            pdf2docx_available = True
        except ImportError:
            pass
        
        try:
            from docx2pdf import convert
            docx2pdf_available = True
        except ImportError:
            pass
        
        try:
            from PIL import Image
            pil_available = True
        except ImportError:
            pass
        
        # 创建转换器实例来获取支持格式
        converter_instance = PDFConverter()
        
        status_info = {
            'pdf_to_word': pdf2docx_available or FITZ_AVAILABLE,
            'word_to_pdf': docx2pdf_available,
            'pdf_to_image': FITZ_AVAILABLE and pil_available,
            'image_to_pdf': pil_available,
            'pdf_to_text': FITZ_AVAILABLE,
            'pdf_processing': FITZ_AVAILABLE,
            'word_processing': pdf2docx_available or docx2pdf_available,
            'image_processing': pil_available,
            'python_version': f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}",
            'server_time': datetime.now().isoformat(),
            'supported_formats': converter_instance.supported_formats
        }
        
        return JsonResponse({
            'success': True,
            'status': 'ready',
            'features': status_info
        })
        
    except Exception as e:
        logger.error(f"PDF转换器状态检查失败: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f'状态检查失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["GET"])
@login_required
def pdf_converter_stats_api(request):
    """获取PDF转换统计数据的API"""
    try:
        from .models.legacy_models import PDFConversionRecord
        from django.db.models import Count, Avg, Sum
        from django.utils import timezone
        from datetime import timedelta
        
        # 获取用户的所有转换记录
        user_records = PDFConversionRecord.objects.filter(user=request.user)
        
        # 总转换次数
        total_conversions = user_records.count()
        
        # 成功转换次数
        successful_conversions = user_records.filter(status='success').count()
        
        # 处理文件数（去重）
        total_files = user_records.values('original_filename').distinct().count()
        
        # 平均转换时间
        avg_conversion_time = user_records.filter(status='success').aggregate(
            avg_time=Avg('conversion_time')
        )['avg_time'] or 0.0
        
        # 用户满意度（基于用户评分）
        rated_conversions = user_records.filter(status='success', satisfaction_rating__isnull=False)
        if rated_conversions.exists():
            avg_rating = rated_conversions.aggregate(avg_rating=Avg('satisfaction_rating'))['avg_rating']
            user_satisfaction = (avg_rating / 5.0) * 100  # 转换为百分比
        else:
            # 如果没有评分记录，使用成功率作为默认值
            user_satisfaction = (successful_conversions / total_conversions * 100) if total_conversions > 0 else 0.0
        
        # 最近转换记录
        recent_conversions = user_records.filter(status='success').order_by('-created_at')[:5]
        recent_data = []
        
        for record in recent_conversions:
            recent_data.append({
                'id': record.id,
                'filename': record.original_filename,
                'conversion_type': record.get_conversion_type_display(),
                'file_size': record.get_file_size_display(),
                'conversion_time': record.get_conversion_time_display(),
                'created_at': record.created_at.strftime('%m-%d %H:%M'),
                'status': record.status,
                'satisfaction_rating': record.satisfaction_rating
            })
        
        return JsonResponse({
            'success': True,
            'stats': {
                'total_conversions': total_conversions,
                'total_files': total_files,
                'avg_speed': f"{avg_conversion_time:.1f}s",
                'user_satisfaction': f"{user_satisfaction:.1f}%"
            },
            'recent_conversions': recent_data
        })
        
    except Exception as e:
        logger.error(f"获取转换统计失败: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def pdf_converter_rating_api(request):
    """更新PDF转换满意度评分的API"""
    try:
        from .models.legacy_models import PDFConversionRecord
        import json
        
        data = json.loads(request.body)
        record_id = data.get('record_id')
        rating = data.get('rating')
        
        if not record_id or not rating:
            return JsonResponse({
                'success': False,
                'error': '缺少必要参数'
            }, status=400)
        
        if not isinstance(rating, int) or rating < 1 or rating > 5:
            return JsonResponse({
                'success': False,
                'error': '评分必须在1-5之间'
            }, status=400)
        
        # 查找转换记录
        try:
            record = PDFConversionRecord.objects.get(id=record_id, user=request.user)
        except PDFConversionRecord.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': '转换记录不存在'
            }, status=404)
        
        # 更新评分
        record.satisfaction_rating = rating
        record.save()
        
        return JsonResponse({
            'success': True,
            'message': '评分更新成功',
            'rating': rating
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': '无效的JSON数据'
        }, status=400)
    except Exception as e:
        logger.error(f"更新满意度评分失败: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f'服务器错误: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def pdf_converter_batch(request):
    """批量PDF转换API"""
    try:
        from .models.legacy_models import PDFConversionRecord
        import time
        
        # 添加调试信息
        logger.info(f"批量PDF转换API请求: POST数据={dict(request.POST)}, FILES数量={len(request.FILES.getlist('files', []))}")
        
        files = request.FILES.getlist('files')
        conversion_type = request.POST.get('type', '')
        
        if not files:
            return JsonResponse({
                'success': False,
                'error': '没有上传文件'
            }, status=400)
        
        if len(files) > 10:
            return JsonResponse({
                'success': False,
                'error': '一次最多只能转换10个文件'
            }, status=400)
        
        results = []
        converter = PDFConverter()
        
        for file in files:
            # 验证文件
            if conversion_type == 'pdf-to-word':
                is_valid, message = converter.validate_file(file, 'pdf')
            elif conversion_type == 'word-to-pdf':
                is_valid, message = converter.validate_file(file, 'word')
            elif conversion_type == 'pdf-to-image':
                is_valid, message = converter.validate_file(file, 'pdf')
            elif conversion_type == 'image-to-pdf':
                is_valid, message = converter.validate_file(file, 'image')
            elif conversion_type == 'pdf-to-text':
                is_valid, message = converter.validate_file(file, 'pdf')
            else:
                is_valid, message = True, "文件验证通过"
            
            if not is_valid:
                # 检查是否包含转换建议
                if "建议的转换类型" in message:
                    # 解析建议的转换类型
                    suggested_types = []
                    lines = message.split('\n')
                    for line in lines:
                        if line.strip().startswith('•'):
                            # 提取转换类型
                            conv_type = line.split('(')[1].split(')')[0] if '(' in line else None
                            if conv_type:
                                suggested_types.append(conv_type)
                    
                    results.append({
                        'filename': file.name,
                        'success': False,
                        'error': message,
                        'suggested_types': suggested_types,
                        'needs_type_switch': True
                    })
                else:
                    results.append({
                        'filename': file.name,
                        'success': False,
                        'error': message
                    })
                continue
            
            # 创建转换记录
            conversion_record = None
            if request.user.is_authenticated:
                conversion_record = PDFConversionRecord.objects.create(
                    user=request.user,
                    conversion_type=conversion_type.replace('-', '_'),
                    original_filename=file.name,
                    file_size=file.size,
                    status='processing'
                )
            
            start_time = time.time()
            
            # 执行转换
            if conversion_type == 'pdf-to-word':
                success, result, file_type = converter.pdf_to_word(file)
            elif conversion_type == 'word-to-pdf':
                success, result, file_type = converter.word_to_pdf(file)
            elif conversion_type == 'pdf-to-image':
                success, result, file_type = converter.pdf_to_images(file)
            elif conversion_type == 'image-to-pdf':
                success, result, file_type = converter.images_to_pdf([file])
            elif conversion_type == 'pdf-to-text':
                success, result, file_type = converter.pdf_to_text(file)
            else:
                success, result, file_type = False, "不支持的转换类型", None
            
            # 计算转换时间
            conversion_time = time.time() - start_time
            
            if success:
                # 处理不同类型的输出
                if file_type == 'pdf_to_images':
                    # 创建ZIP文件包含所有图片
                    import zipfile
                    from io import BytesIO
                    import base64
                    
                    zip_buffer = BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                        for i, img_data in enumerate(result):
                            # 解码base64图片数据
                            img_bytes = base64.b64decode(img_data['data'])
                            # 添加到ZIP文件
                            zip_file.writestr(f'{file.name}_page_{i+1}.png', img_bytes)
                    
                    zip_content = zip_buffer.getvalue()
                    zip_buffer.close()
                    
                    # 保存ZIP文件
                    output_filename = f"{uuid.uuid4()}_{conversion_type.replace('-', '_')}_images.zip"
                    file_path = default_storage.save(f'converted/{output_filename}', ContentFile(zip_content))
                    download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
                    
                    # 更新转换记录
                    if conversion_record:
                        conversion_record.status = 'success'
                        conversion_record.output_filename = output_filename
                        conversion_record.conversion_time = conversion_time
                        conversion_record.download_url = download_url
                        conversion_record.save()
                    
                    results.append({
                        'filename': file.name,
                        'success': True,
                        'download_url': download_url,
                        'output_filename': output_filename,
                        'total_pages': len(result),
                        'message': f'已转换{len(result)}页，打包为ZIP文件供下载'
                    })
                    continue
                
                # 保存文件
                output_filename = f"{uuid.uuid4()}_{conversion_type.replace('-', '_')}"
                if file_type == 'pdf_to_word':
                    output_filename += '.docx'
                elif file_type == 'word_to_pdf':
                    output_filename += '.pdf'
                elif file_type == 'images_to_pdf':
                    output_filename += '.pdf'
                elif file_type == 'pdf_to_text':
                    output_filename += '.txt'
                    result = result.encode('utf-8')
                
                file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
                download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
                
                # 更新转换记录
                if conversion_record:
                    conversion_record.status = 'success'
                    conversion_record.output_filename = output_filename
                    conversion_record.conversion_time = conversion_time
                    conversion_record.download_url = download_url
                    conversion_record.save()
                
                results.append({
                    'filename': file.name,
                    'success': True,
                    'download_url': download_url,
                    'output_filename': output_filename
                })
            else:
                # 更新转换记录为失败状态
                if conversion_record:
                    conversion_record.status = 'failed'
                    conversion_record.error_message = result
                    conversion_record.conversion_time = conversion_time
                    conversion_record.save()
                
                results.append({
                    'filename': file.name,
                    'success': False,
                    'error': result
                })
        
        return JsonResponse({
            'success': True,
            'results': results,
            'total_files': len(files),
            'successful_conversions': len([r for r in results if r['success']])
        })
        
    except Exception as e:
        logger.error(f"批量PDF转换API错误: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f'服务器错误: {str(e)}'
        }, status=500) 

@csrf_exempt
@require_http_methods(["GET"])
def pdf_download_view(request, filename):
    """
    专门的PDF文件下载视图，解决Google浏览器下载问题
    """
    try:
        from django.http import FileResponse, Http404, HttpResponse
        from django.conf import settings
        import os
        import mimetypes
        
        # 构建文件路径
        file_path = os.path.join(settings.MEDIA_ROOT, 'converted', filename)
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise Http404("文件不存在")
        
        # 获取文件大小
        file_size = os.path.getsize(file_path)
        
        # 确定MIME类型
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain',
            '.zip': 'application/zip',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.tiff': 'image/tiff'
        }
        
        file_ext = os.path.splitext(filename)[1].lower()
        content_type = mime_types.get(file_ext, 'application/octet-stream')
        
        # 如果MIME类型未知，尝试自动检测
        if content_type == 'application/octet-stream':
            detected_type, _ = mimetypes.guess_type(filename)
            if detected_type:
                content_type = detected_type
        
        logger.info(f"下载文件: {filename}, 路径: {file_path}, 大小: {file_size}, 类型: {content_type}")
        
        # 打开文件并创建响应
        try:
            file_handle = open(file_path, 'rb')
            response = FileResponse(file_handle, content_type=content_type)
        except Exception as e:
            logger.error(f"打开文件失败: {str(e)}")
            raise Http404("文件读取失败")
        
        # 设置下载头信息
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = file_size
        
        # 添加缓存控制头，防止浏览器缓存
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        
        # 添加CORS头，允许跨域下载
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Content-Disposition'
        
        # 添加额外的下载头
        response['X-Content-Type-Options'] = 'nosniff'
        response['X-Frame-Options'] = 'DENY'
        
        logger.info(f"文件下载响应已创建: {filename}")
        return response
            
    except Http404 as e:
        logger.error(f"文件不存在: {filename}")
        return HttpResponse(f"文件不存在: {filename}", status=404)
    except Exception as e:
        logger.error(f"PDF下载视图错误: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f'下载失败: {str(e)}'
        }, status=500) 

@csrf_exempt
@require_http_methods(["POST"])
def pdf_converter_test_api(request):
    """PDF转换测试API（无需登录）"""
    try:
        import time
        
        # 添加调试信息
        logger.info(f"PDF转换测试API请求: POST数据={dict(request.POST)}, FILES={list(request.FILES.keys())}")
        
        conversion_type = request.POST.get('type', '')
        
        # 检查是否有文件上传（文本转PDF除外）
        if conversion_type != 'text-to-pdf':
            if 'file' not in request.FILES:
                logger.warning("PDF转换测试API: 没有上传文件")
                return JsonResponse({
                    'success': False,
                    'error': '没有上传文件'
                }, status=400)
            file = request.FILES['file']
        else:
            # 文本转PDF不需要文件上传
            file = None
        
        start_time = time.time()
        
        # 验证转换类型
        valid_types = ['pdf-to-word', 'word-to-pdf', 'pdf-to-image', 'image-to-pdf', 'text-to-pdf', 'pdf-to-text', 'txt-to-pdf']
        if conversion_type not in valid_types:
            return JsonResponse({
                'success': False,
                'error': f'不支持的转换类型: {conversion_type}'
            }, status=400)
        
        # 根据转换类型验证文件格式
        if conversion_type == 'text-to-pdf':
            # 文本转PDF不需要文件验证
            text_content = request.POST.get('text_content', '')
            if not text_content.strip():
                is_valid, message = False, "请输入要转换的文本内容"
            else:
                is_valid, message = True, "文本内容验证通过"
        elif conversion_type == 'pdf-to-word':
            is_valid, message = converter.validate_file(file, 'pdf')
        elif conversion_type == 'word-to-pdf':
            is_valid, message = converter.validate_file(file, 'word')
        elif conversion_type == 'pdf-to-image':
            is_valid, message = converter.validate_file(file, 'pdf')
        elif conversion_type == 'image-to-pdf':
            is_valid, message = converter.validate_file(file, 'image')
        elif conversion_type == 'pdf-to-text':
            is_valid, message = converter.validate_file(file, 'pdf')
        elif conversion_type == 'txt-to-pdf':
            is_valid, message = converter.validate_file(file, 'text')
        else:
            is_valid, message = True, "文件验证通过"
        
        if not is_valid:
            return JsonResponse({
                'success': False,
                'error': message
            }, status=400)
        
        # 执行转换
        if conversion_type == 'pdf-to-word':
            success, result, file_type = converter.pdf_to_word(file)
        elif conversion_type == 'word-to-pdf':
            success, result, file_type = converter.word_to_pdf(file)
        elif conversion_type == 'pdf-to-image':
            success, result, file_type = converter.pdf_to_images(file)
        elif conversion_type == 'image-to-pdf':
            success, result, file_type = converter.images_to_pdf([file])
        elif conversion_type == 'pdf-to-text':
            success, result, file_type = converter.pdf_to_text(file)
        elif conversion_type == 'text-to-pdf':
            success, result, file_type = converter.text_to_pdf(request.POST.get('text_content', ''))
        elif conversion_type == 'txt-to-pdf':
            success, result, file_type = converter.txt_to_pdf(file)
        else:
            return JsonResponse({
                'success': False,
                'error': f'不支持的转换类型: {conversion_type}'
            }, status=400)
        
        if not success:
            return JsonResponse({
                'success': False,
                'error': result
            }, status=500)
        
        # 保存转换结果
        output_filename = f"{uuid.uuid4()}_{conversion_type.replace('-', '_')}"
        
        if file_type == 'pdf_to_word':
            output_filename += '.docx'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        elif file_type == 'word_to_pdf':
            output_filename += '.pdf'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        elif file_type == 'pdf_to_images':
            # 创建ZIP文件包含所有图片
            import zipfile
            from io import BytesIO
            import base64
            
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for i, img_data in enumerate(result):
                    # 解码base64图片数据
                    img_bytes = base64.b64decode(img_data['data'])
                    # 添加到ZIP文件
                    zip_file.writestr(f'page_{i+1}.png', img_bytes)
            
            zip_content = zip_buffer.getvalue()
            zip_buffer.close()
            
            # 保存ZIP文件
            output_filename += '_images.zip'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(zip_content))
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
            
            # 返回下载链接
            return JsonResponse({
                'success': True,
                'type': 'file',
                'download_url': download_url,
                'filename': output_filename,
                'original_filename': file.name if file else '图片集合',
                'file_size': len(zip_content),
                'total_pages': len(result),
                'message': f'已转换{len(result)}页，打包为ZIP文件供下载',
                'conversion_type': conversion_type
            })
        elif file_type == 'images_to_pdf':
            output_filename += '.pdf'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        elif file_type == 'text_to_pdf':
            output_filename += '.pdf'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        elif file_type == 'pdf_to_text':
            output_filename += '.txt'
            file_path = default_storage.save(f'converted/{output_filename}', ContentFile(result.encode('utf-8')))
            # 设置下载链接
            download_url = f'/tools/api/pdf-converter/download/{output_filename}/'
        else:
            return JsonResponse({
                'success': False,
                'error': '未知的文件类型'
            }, status=500)
        
        # 确定原始文件名
        if conversion_type == 'text-to-pdf':
            original_filename = '文本内容'
        elif conversion_type == 'pdf-to-text':
            original_filename = file.name if file else 'PDF文件'
        else:
            original_filename = file.name if file else '文件'
        
        return JsonResponse({
            'success': True,
            'type': 'file',
            'download_url': download_url,
            'filename': output_filename,
            'original_filename': original_filename,
            'conversion_type': conversion_type
        })
        
    except Exception as e:
        logger.error(f"PDF转换测试API错误: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f'服务器错误: {str(e)}'
        }, status=500) 