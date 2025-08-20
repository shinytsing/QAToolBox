#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试Word转PDF的中文字符显示
"""

import os
import sys
import django

# 设置Django环境
sys.path.append('/Users/gaojie/PycharmProjects/QAToolBox')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')
django.setup()

from apps.tools.pdf_converter_api import PDFConverter
from django.core.files.uploadedfile import SimpleUploadedFile
import io

def test_chinese_word_to_pdf():
    """测试包含中文的Word转PDF"""
    print("🔍 测试中文Word转PDF...")
    converter = PDFConverter()
    
    try:
        # 创建包含中文的Word文档
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('中文测试文档', 0)
        
        # 添加段落
        doc.add_paragraph('这是一个包含中文内容的测试文档。')
        doc.add_paragraph('测试内容包括：')
        
        # 添加列表
        items = [
            '中文字符显示',
            '特殊符号：！@#￥%……&*（）',
            '数字：1234567890',
            '英文：Hello World',
            '混合内容：中文123ABC！@#'
        ]
        
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
        
        # 添加更多中文内容
        doc.add_paragraph('')
        doc.add_paragraph('这是一个较长的段落，用来测试中文字符在PDF中的显示效果。包含各种标点符号和特殊字符，确保转换后的PDF能够正确显示所有内容。')
        
        doc.add_paragraph('测试完成！')
        
        # 保存到内存
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        word_content = doc_buffer.getvalue()
        word_file = SimpleUploadedFile("chinese_test.docx", word_content, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        print("   正在转换包含中文的Word文档...")
        success, result, file_type = converter.word_to_pdf(word_file)
        
        if success:
            print("✅ 中文Word转PDF成功")
            print(f"   生成PDF大小: {len(result)} bytes")
            
            # 保存PDF文件用于检查
            output_path = "test_chinese_output.pdf"
            with open(output_path, 'wb') as f:
                f.write(result)
            print(f"   PDF已保存到: {output_path}")
            print("   请打开PDF文件检查中文字符是否正确显示")
            
        else:
            print(f"❌ 中文Word转PDF失败: {result}")
            
        word_file.close()
        doc_buffer.close()
        
    except ImportError:
        print("⚠️  python-docx库未安装，跳过中文Word转PDF测试")
    except Exception as e:
        print(f"❌ 中文Word转PDF异常: {str(e)}")

if __name__ == "__main__":
    print("🚀 开始测试中文Word转PDF...")
    print("=" * 50)
    
    test_chinese_word_to_pdf()
    
    print("=" * 50)
    print("🎉 测试完成！")
