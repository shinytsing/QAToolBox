#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
详细的docx2pdf测试脚本
"""

import os
import sys
import tempfile
import subprocess
import io
from pathlib import Path

def test_docx2pdf_direct():
    """直接测试docx2pdf"""
    print("🧪 直接测试docx2pdf...")
    
    try:
        from docx2pdf import convert
        from docx import Document
        
        # 创建测试文档
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test Word document.')
        doc.add_paragraph('Testing PDF conversion.')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        try:
            print(f"   输入文件: {temp_docx_path}")
            print(f"   输出文件: {temp_pdf_path}")
            
            # 检查输入文件
            if os.path.exists(temp_docx_path):
                print(f"   ✅ 输入文件存在，大小: {os.path.getsize(temp_docx_path)} 字节")
            else:
                print("   ❌ 输入文件不存在")
                return False
            
            # 执行转换
            print("   开始转换...")
            convert(temp_docx_path, temp_pdf_path)
            
            # 检查输出文件
            if os.path.exists(temp_pdf_path):
                size = os.path.getsize(temp_pdf_path)
                print(f"   ✅ 输出文件存在，大小: {size} 字节")
                if size > 0:
                    print("   ✅ 转换成功!")
                    return True
                else:
                    print("   ❌ 输出文件为空")
                    return False
            else:
                print("   ❌ 输出文件不存在")
                return False
                
        except Exception as e:
            print(f"   ❌ 转换失败: {e}")
            import traceback
            traceback.print_exc()
            return False
        finally:
            # 清理文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
                
    except ImportError as e:
        print(f"❌ docx2pdf导入失败: {e}")
        return False

def test_microsoft_word():
    """检查Microsoft Word是否可用"""
    print("\n🔍 检查Microsoft Word...")
    
    word_paths = [
        "/Applications/Microsoft Word.app",
        "/Applications/Microsoft Office 2019/Microsoft Word.app",
        "/Applications/Microsoft Office 2021/Microsoft Word.app"
    ]
    
    for path in word_paths:
        if os.path.exists(path):
            print(f"✅ 找到Microsoft Word: {path}")
            return True
    
    print("❌ 未找到Microsoft Word")
    return False

def test_alternative_conversion():
    """测试替代转换方法"""
    print("\n🔧 测试替代转换方法...")
    
    try:
        from docx import Document
        
        # 创建测试文档
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test Word document.')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        try:
            # 尝试使用Pandoc
            print("   尝试使用Pandoc...")
            result = subprocess.run([
                'pandoc', temp_docx_path, '-o', temp_pdf_path, '--pdf-engine=wkhtmltopdf'
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0 and os.path.exists(temp_pdf_path):
                size = os.path.getsize(temp_pdf_path)
                print(f"   ✅ Pandoc转换成功，文件大小: {size} 字节")
                return True
            else:
                print(f"   ❌ Pandoc转换失败: {result.stderr}")
                
                # 尝试使用LibreOffice
                print("   尝试使用LibreOffice...")
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(temp_pdf_path), temp_docx_path
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0 and os.path.exists(temp_pdf_path):
                    size = os.path.getsize(temp_pdf_path)
                    print(f"   ✅ LibreOffice转换成功，文件大小: {size} 字节")
                    return True
                else:
                    print(f"   ❌ LibreOffice转换失败: {result.stderr}")
                    return False
                    
        finally:
            # 清理文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
                
    except Exception as e:
        print(f"❌ 替代转换测试失败: {e}")
        return False

def main():
    """主函数"""
    print("🔍 docx2pdf详细诊断")
    print("=" * 50)
    
    # 检查Microsoft Word
    word_available = test_microsoft_word()
    
    # 测试docx2pdf
    docx2pdf_success = test_docx2pdf_direct()
    
    # 测试替代方法
    alternative_success = test_alternative_conversion()
    
    print("\n" + "=" * 50)
    print("诊断结果:")
    
    if docx2pdf_success:
        print("✅ docx2pdf工作正常")
    else:
        print("❌ docx2pdf有问题")
        if not word_available:
            print("   原因: 未找到Microsoft Word")
            print("   解决方案: 安装Microsoft Word或使用替代方法")
    
    if alternative_success:
        print("✅ 替代转换方法可用")
    else:
        print("❌ 替代转换方法不可用")
    
    if docx2pdf_success or alternative_success:
        print("\n✅ Word转PDF功能可用")
    else:
        print("\n❌ Word转PDF功能不可用")
        print("建议安装Microsoft Word或LibreOffice")

if __name__ == "__main__":
    main() 