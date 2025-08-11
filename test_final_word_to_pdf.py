#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最终的Word转PDF测试脚本
"""

import os
import sys
import tempfile
import io

# 设置Django环境
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')

def test_current_word_to_pdf():
    """测试当前的Word转PDF功能"""
    print("🧪 测试当前的Word转PDF功能...")
    
    try:
        import django
        django.setup()
        
        from apps.tools.pdf_converter_api import PDFConverter
        from docx import Document
        
        # 创建测试文档
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test Word document.')
        doc.add_paragraph('Testing current PDF conversion.')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        try:
            # 创建文件对象
            with open(temp_docx_path, 'rb') as f:
                word_file = io.BytesIO(f.read())
                word_file.name = "test.docx"
            
            # 测试转换
            converter = PDFConverter()
            success, result, file_type = converter.word_to_pdf(word_file)
            
            if success:
                print("✅ 当前Word转PDF转换成功!")
                print(f"   输出类型: {file_type}")
                print(f"   输出大小: {len(result)} 字节")
                
                # 保存测试结果
                with open('test_current_output.pdf', 'wb') as f:
                    f.write(result)
                print("   测试结果已保存到: test_current_output.pdf")
                return True
            else:
                print(f"❌ 当前Word转PDF转换失败: {result}")
                return False
                
        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
                
    except Exception as e:
        print(f"❌ 当前Word转PDF测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_pandoc_direct():
    """直接测试Pandoc"""
    print("\n🧪 直接测试Pandoc...")
    
    try:
        from docx import Document
        import subprocess
        
        # 创建测试文档
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test Word document.')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        try:
            print(f"   输入文件: {temp_docx_path}")
            print(f"   输出文件: {temp_pdf_path}")
            
            # 使用Pandoc转换
            result = subprocess.run([
                'pandoc', temp_docx_path, '-o', temp_pdf_path
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and os.path.exists(temp_pdf_path):
                size = os.path.getsize(temp_pdf_path)
                print(f"   ✅ Pandoc转换成功，文件大小: {size} 字节")
                
                # 保存结果
                with open(temp_pdf_path, 'rb') as f:
                    content = f.read()
                with open('test_pandoc_output.pdf', 'wb') as f:
                    f.write(content)
                print("   结果已保存到: test_pandoc_output.pdf")
                return True
            else:
                print(f"   ❌ Pandoc转换失败: {result.stderr}")
                return False
                
        finally:
            # 清理文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
                
    except Exception as e:
        print(f"❌ Pandoc测试失败: {e}")
        return False

def main():
    """主函数"""
    print("🔍 最终Word转PDF测试")
    print("=" * 50)
    
    # 测试当前功能
    current_success = test_current_word_to_pdf()
    
    # 测试Pandoc
    pandoc_success = test_pandoc_direct()
    
    print("\n" + "=" * 50)
    print("测试结果:")
    
    if current_success:
        print("✅ 当前Word转PDF功能正常")
    else:
        print("❌ 当前Word转PDF功能有问题")
    
    if pandoc_success:
        print("✅ Pandoc转换功能正常")
    else:
        print("❌ Pandoc转换功能有问题")
    
    if current_success or pandoc_success:
        print("\n✅ Word转PDF功能可用")
        if not current_success and pandoc_success:
            print("建议: 可以集成Pandoc作为Word转PDF的解决方案")
    else:
        print("\n❌ Word转PDF功能不可用")
        print("建议: 安装Microsoft Word或LibreOffice")

if __name__ == "__main__":
    main() 