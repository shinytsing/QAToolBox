#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Pandoc Word转PDF测试脚本
"""

import os
import sys
import tempfile
import subprocess
import io
from pathlib import Path

def test_pandoc_direct():
    """直接测试Pandoc转换"""
    print("🧪 直接测试Pandoc转换...")
    
    try:
        from docx import Document
        
        # 创建一个简单的测试Word文档
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个测试Word文档。')
        doc.add_paragraph('包含中文和English内容。')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        try:
            print(f"   临时docx文件: {temp_docx_path}")
            print(f"   临时pdf文件: {temp_pdf_path}")
            
            # 使用Pandoc转换
            result = subprocess.run([
                'pandoc', temp_docx_path, '-o', temp_pdf_path
            ], capture_output=True, text=True, timeout=30)
            
            print(f"   Pandoc返回码: {result.returncode}")
            if result.stdout:
                print(f"   输出: {result.stdout}")
            if result.stderr:
                print(f"   错误: {result.stderr}")
            
            # 检查输出文件
            if os.path.exists(temp_pdf_path) and os.path.getsize(temp_pdf_path) > 0:
                print("✅ Pandoc转换测试成功")
                print(f"   输出文件大小: {os.path.getsize(temp_pdf_path)} 字节")
                return True
            else:
                print("❌ Pandoc转换测试失败")
                return False
                
        except Exception as e:
            print(f"❌ Pandoc转换测试失败: {e}")
            import traceback
            traceback.print_exc()
            return False
        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
                
    except ImportError as e:
        print(f"❌ python-docx导入失败: {e}")
        return False

def test_pdf_converter_with_pandoc():
    """测试PDF转换器的Pandoc方法"""
    print("\n🧪 测试PDF转换器的Pandoc方法...")
    
    try:
        # 设置Django环境
        os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')
        
        import django
        django.setup()
        
        from apps.tools.pdf_converter_api import PDFConverter
        from docx import Document
        
        # 创建一个简单的测试Word文档
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个测试Word文档。')
        doc.add_paragraph('包含中文和English内容。')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        try:
            # 创建文件对象
            with open(temp_docx_path, 'rb') as f:
                word_file = io.BytesIO(f.read())
                word_file.name = "test.docx"
            
            # 测试转换
            converter = PDFConverter()
            success, result, file_type = converter.word_to_pdf_alternative(word_file)
            
            if success:
                print("✅ PDF转换器Pandoc方法成功!")
                print(f"   输出类型: {file_type}")
                print(f"   输出大小: {len(result)} 字节")
                return True
            else:
                print(f"❌ PDF转换器Pandoc方法失败: {result}")
                return False
                
        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
                
    except Exception as e:
        print(f"❌ PDF转换器测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """主函数"""
    print("🔍 Pandoc Word转PDF测试")
    print("=" * 50)
    
    # 测试Pandoc直接转换
    pandoc_success = test_pandoc_direct()
    
    # 测试PDF转换器
    converter_success = test_pdf_converter_with_pandoc()
    
    print("\n" + "=" * 50)
    if pandoc_success and converter_success:
        print("✅ Pandoc Word转PDF功能完全正常!")
        print("现在可以使用Word转PDF功能了")
    elif pandoc_success:
        print("⚠️  Pandoc直接转换正常，但PDF转换器有问题")
    else:
        print("❌ Pandoc转换失败")
        print("请检查Pandoc安装和配置")

if __name__ == "__main__":
    main() 