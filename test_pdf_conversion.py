#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF转换功能测试脚本
验证pdf2docx和docx2pdf库是否正常工作
"""

import os
import sys
import tempfile
import io
from pathlib import Path

# 设置Django环境
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')

def test_pdf2docx():
    """测试pdf2docx库"""
    print("🧪 测试pdf2docx库...")
    
    try:
        from pdf2docx import Converter
        print("✅ pdf2docx导入成功")
        
        # 创建一个简单的测试PDF
        test_pdf_content = b'%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/MediaBox [0 0 612 792]\n/Contents 4 0 R\n>>\nendobj\n4 0 obj\n<<\n/Length 44\n>>\nstream\nBT\n/F1 12 Tf\n72 720 Td\n(Test PDF Content) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000204 00000 n \ntrailer\n<<\n/Size 5\n/Root 1 0 R\n>>\nstartxref\n297\n%%EOF'
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf.write(test_pdf_content)
            temp_pdf_path = temp_pdf.name
        
        temp_docx_path = temp_pdf_path.replace('.pdf', '.docx')
        
        try:
            # 测试转换
            cv = Converter(temp_pdf_path)
            cv.convert(temp_docx_path)
            cv.close()
            
            # 检查输出文件
            if os.path.exists(temp_docx_path) and os.path.getsize(temp_docx_path) > 0:
                print("✅ pdf2docx转换测试成功")
                print(f"   输出文件大小: {os.path.getsize(temp_docx_path)} 字节")
            else:
                print("❌ pdf2docx转换测试失败")
                
        except Exception as e:
            print(f"❌ pdf2docx转换测试失败: {e}")
        finally:
            # 清理临时文件
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
                
    except ImportError as e:
        print(f"❌ pdf2docx导入失败: {e}")

def test_docx2pdf():
    """测试docx2pdf库"""
    print("\n🧪 测试docx2pdf库...")
    
    try:
        from docx2pdf import convert
        print("✅ docx2pdf导入成功")
        
        # 创建一个简单的测试Word文档
        from docx import Document
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个测试Word文档。')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        try:
            # 测试转换
            convert(temp_docx_path, temp_pdf_path)
            
            # 检查输出文件
            if os.path.exists(temp_pdf_path) and os.path.getsize(temp_pdf_path) > 0:
                print("✅ docx2pdf转换测试成功")
                print(f"   输出文件大小: {os.path.getsize(temp_pdf_path)} 字节")
            else:
                print("❌ docx2pdf转换测试失败")
                
        except Exception as e:
            print(f"❌ docx2pdf转换测试失败: {e}")
        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
                
    except ImportError as e:
        print(f"❌ docx2pdf导入失败: {e}")

def test_django_pdf_converter():
    """测试Django PDF转换器"""
    print("\n🧪 测试Django PDF转换器...")
    
    try:
        import django
        django.setup()
        
        from apps.tools.pdf_converter_api import PDFConverter
        converter = PDFConverter()
        print("✅ Django PDF转换器导入成功")
        
        # 检查pdf2docx可用性
        if hasattr(converter, 'PDF2DOCX_AVAILABLE'):
            if converter.PDF2DOCX_AVAILABLE:
                print("✅ pdf2docx在PDF转换器中可用")
            else:
                print("❌ pdf2docx在PDF转换器中不可用")
        
        # 检查docx2pdf可用性
        if hasattr(converter, 'DOCX2PDF_AVAILABLE'):
            if converter.DOCX2PDF_AVAILABLE:
                print("✅ docx2pdf在PDF转换器中可用")
            else:
                print("❌ docx2pdf在PDF转换器中不可用")
                
    except Exception as e:
        print(f"❌ Django PDF转换器测试失败: {e}")

def main():
    """主函数"""
    print("🔍 PDF转换功能测试")
    print("=" * 50)
    
    # 测试各个库
    test_pdf2docx()
    test_docx2pdf()
    test_django_pdf_converter()
    
    print("\n✅ 测试完成！")
    print("如果所有测试都通过，PDF转换功能应该可以正常使用。")

if __name__ == "__main__":
    main() 