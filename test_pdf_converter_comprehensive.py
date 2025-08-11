#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF转换引擎全面测试脚本
测试所有转换功能并验证下载功能
"""

import os
import sys
import requests
import tempfile
import time
import json
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont

# 添加项目路径
sys.path.append('/Users/gaojie/PycharmProjects/QAToolBox')

def create_test_pdf():
    """创建测试PDF文件"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # 创建临时PDF文件
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_pdf_path = temp_file.name
        
        # 创建PDF内容
        c = canvas.Canvas(temp_pdf_path, pagesize=letter)
        c.drawString(100, 750, "PDF转换引擎测试文档")
        c.drawString(100, 720, "这是一个用于测试PDF转换功能的文档。")
        c.drawString(100, 690, "包含中文和英文内容。")
        c.drawString(100, 660, "Test PDF Conversion Engine")
        c.drawString(100, 630, "This is a test document for PDF conversion.")
        c.drawString(100, 600, "Contains both Chinese and English content.")
        c.save()
        
        return temp_pdf_path
    except Exception as e:
        print(f"创建测试PDF失败: {e}")
        return None

def create_test_image():
    """创建测试图片文件"""
    try:
        # 创建测试图片
        img = Image.new('RGB', (400, 300), color='white')
        draw = ImageDraw.Draw(img)
        
        # 添加文字
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 20)
        except:
            font = ImageFont.load_default()
        
        draw.text((50, 50), "测试图片", fill='black', font=font)
        draw.text((50, 100), "Test Image", fill='black', font=font)
        draw.text((50, 150), "用于PDF转换测试", fill='black', font=font)
        
        # 保存为临时文件
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
            temp_img_path = temp_file.name
        
        img.save(temp_img_path, 'PNG')
        return temp_img_path
    except Exception as e:
        print(f"创建测试图片失败: {e}")
        return None

def create_test_word():
    """创建测试Word文档"""
    try:
        from docx import Document
        
        # 创建Word文档
        doc = Document()
        doc.add_heading('Word转PDF测试文档', 0)
        doc.add_paragraph('这是一个用于测试Word转PDF功能的文档。')
        doc.add_paragraph('包含中文和英文内容。')
        doc.add_paragraph('Test Word to PDF Conversion')
        doc.add_paragraph('This is a test document for Word to PDF conversion.')
        
        # 保存为临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_docx_path = temp_file.name
        
        doc.save(temp_docx_path)
        return temp_docx_path
    except Exception as e:
        print(f"创建测试Word文档失败: {e}")
        return None

def test_pdf_conversion_api(conversion_type, file_path=None, text_content=None):
    """测试PDF转换API"""
    url = 'http://localhost:8000/tools/api/pdf-converter/'
    
    data = {'type': conversion_type}
    files = {}
    
    if file_path and os.path.exists(file_path):
        with open(file_path, 'rb') as f:
            files['file'] = (os.path.basename(file_path), f.read())
    
    if text_content:
        data['text_content'] = text_content
    
    try:
        print(f"🧪 测试 {conversion_type}...")
        response = requests.post(url, data=data, files=files, timeout=30)
        
        if response.status_code == 200:
            result = response.json()
            if result.get('success'):
                print(f"✅ {conversion_type} 转换成功!")
                print(f"   下载链接: {result.get('download_url')}")
                print(f"   文件名: {result.get('filename')}")
                
                # 测试下载功能
                download_url = result.get('download_url')
                if download_url:
                    download_test_url = f"http://localhost:8000{download_url}"
                    print(f"   测试下载: {download_test_url}")
                    
                    download_response = requests.get(download_test_url, timeout=30)
                    if download_response.status_code == 200:
                        print(f"✅ 下载测试成功! 文件大小: {len(download_response.content)} 字节")
                        return True, result
                    else:
                        print(f"❌ 下载测试失败: {download_response.status_code}")
                        return False, result
                else:
                    print("⚠️ 没有下载链接")
                    return True, result
            else:
                print(f"❌ {conversion_type} 转换失败: {result.get('error')}")
                return False, result
        else:
            print(f"❌ API请求失败: {response.status_code}")
            print(f"   响应内容: {response.text}")
            return False, None
    except Exception as e:
        print(f"❌ 测试 {conversion_type} 时出错: {e}")
        return False, None

def test_all_conversions():
    """测试所有转换功能"""
    print("🚀 PDF转换引擎全面测试开始")
    print("=" * 60)
    
    # 创建测试文件
    print("📁 创建测试文件...")
    test_pdf = create_test_pdf()
    test_image = create_test_image()
    test_word = create_test_word()
    
    if not test_pdf:
        print("❌ 无法创建测试PDF文件，测试终止")
        return
    
    results = {}
    
    # 1. 测试PDF转Word
    print("\n1️⃣ 测试PDF转Word功能")
    success, result = test_pdf_conversion_api('pdf-to-word', test_pdf)
    results['pdf-to-word'] = success
    
    # 2. 测试Word转PDF
    print("\n2️⃣ 测试Word转PDF功能")
    if test_word:
        success, result = test_pdf_conversion_api('word-to-pdf', test_word)
        results['word-to-pdf'] = success
    else:
        print("⚠️ 跳过Word转PDF测试（无法创建Word文档）")
        results['word-to-pdf'] = False
    
    # 3. 测试PDF转图片
    print("\n3️⃣ 测试PDF转图片功能")
    success, result = test_pdf_conversion_api('pdf-to-image', test_pdf)
    results['pdf-to-image'] = success
    
    # 4. 测试图片转PDF
    print("\n4️⃣ 测试图片转PDF功能")
    if test_image:
        success, result = test_pdf_conversion_api('image-to-pdf', test_image)
        results['image-to-pdf'] = success
    else:
        print("⚠️ 跳过图片转PDF测试（无法创建图片文件）")
        results['image-to-pdf'] = False
    
    # 5. 测试PDF转文本
    print("\n5️⃣ 测试PDF转文本功能")
    success, result = test_pdf_conversion_api('pdf-to-text', test_pdf)
    results['pdf-to-text'] = success
    
    # 6. 测试文本转PDF
    print("\n6️⃣ 测试文本转PDF功能")
    test_text = """这是一个测试文本内容。
This is a test text content.
包含中文和英文。
Contains both Chinese and English.
用于测试文本转PDF功能。
For testing text to PDF conversion."""
    success, result = test_pdf_conversion_api('text-to-pdf', text_content=test_text)
    results['text-to-pdf'] = success
    
    # 7. 测试TXT文件转PDF
    print("\n7️⃣ 测试TXT文件转PDF功能")
    if test_text:
        # 创建临时TXT文件
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False, mode='w', encoding='utf-8') as temp_file:
            temp_file.write(test_text)
            temp_txt_path = temp_file.name
        
        success, result = test_pdf_conversion_api('txt-to-pdf', temp_txt_path)
        results['txt-to-pdf'] = success
        
        # 清理临时文件
        try:
            os.unlink(temp_txt_path)
        except:
            pass
    else:
        print("⚠️ 跳过TXT转PDF测试")
        results['txt-to-pdf'] = False
    
    # 清理临时文件
    print("\n🧹 清理临时文件...")
    for file_path in [test_pdf, test_image, test_word]:
        if file_path and os.path.exists(file_path):
            try:
                os.unlink(file_path)
            except:
                pass
    
    # 输出测试结果总结
    print("\n" + "=" * 60)
    print("📊 测试结果总结")
    print("=" * 60)
    
    total_tests = len(results)
    successful_tests = sum(results.values())
    
    for test_name, success in results.items():
        status = "✅ 成功" if success else "❌ 失败"
        print(f"{test_name:15} : {status}")
    
    print(f"\n总计: {successful_tests}/{total_tests} 个功能测试通过")
    
    if successful_tests == total_tests:
        print("🎉 所有PDF转换功能测试通过！")
    else:
        print("⚠️ 部分功能测试失败，请检查相关依赖和配置")
    
    return results

def test_api_status():
    """测试API状态"""
    print("\n🔍 测试API状态...")
    try:
        response = requests.get('http://localhost:8000/tools/api/pdf-converter/status/', timeout=10)
        if response.status_code == 200:
            result = response.json()
            if result.get('success'):
                features = result.get('features', {})
                print("✅ API状态正常")
                print("📋 功能支持情况:")
                for feature, supported in features.items():
                    if isinstance(supported, bool):
                        status = "✅" if supported else "❌"
                        print(f"   {feature:20} : {status}")
                return True
            else:
                print(f"❌ API状态异常: {result.get('error')}")
                return False
        else:
            print(f"❌ API状态请求失败: {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ API状态测试失败: {e}")
        return False

if __name__ == "__main__":
    print("🚀 PDF转换引擎全面测试")
    print("=" * 60)
    
    # 检查服务器是否运行
    try:
        response = requests.get('http://localhost:8000/', timeout=5)
        print("✅ Django服务器正在运行")
    except:
        print("❌ Django服务器未运行，请先启动服务器")
        print("   命令: python manage.py runserver 0.0.0.0:8000")
        sys.exit(1)
    
    # 测试API状态
    test_api_status()
    
    # 测试所有转换功能
    results = test_all_conversions()
    
    print("\n🎯 测试完成！")
    print("=" * 60)
