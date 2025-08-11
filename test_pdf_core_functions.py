#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF转换引擎核心功能测试
"""

import os
import sys
import tempfile
from io import BytesIO

# 添加项目路径
sys.path.append('/Users/gaojie/PycharmProjects/QAToolBox')

def test_pdf_converter_import():
    """测试PDF转换器导入"""
    try:
        from apps.tools.pdf_converter_api import PDFConverter
        print("✅ PDF转换器导入成功")
        return True
    except Exception as e:
        print(f"❌ PDF转换器导入失败: {e}")
        return False

def test_pdf_converter_initialization():
    """测试PDF转换器初始化"""
    try:
        from apps.tools.pdf_converter_api import PDFConverter
        converter = PDFConverter()
        print("✅ PDF转换器初始化成功")
        print(f"   支持格式: {converter.supported_formats}")
        return True
    except Exception as e:
        print(f"❌ PDF转换器初始化失败: {e}")
        return False

def test_file_validation():
    """测试文件验证功能"""
    try:
        from apps.tools.pdf_converter_api import PDFConverter
        converter = PDFConverter()
        
        # 创建测试文件
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(b'%PDF-1.4\n%Test PDF content')
            temp_pdf_path = temp_file.name
        
        # 模拟文件对象
        class MockFile:
            def __init__(self, path):
                self.name = os.path.basename(path)
                self.size = os.path.getsize(path)
                self.path = path
            
            def read(self):
                with open(self.path, 'rb') as f:
                    return f.read()
        
        mock_file = MockFile(temp_pdf_path)
        
        # 测试PDF文件验证
        is_valid, message = converter.validate_file(mock_file, 'pdf')
        print(f"✅ PDF文件验证: {is_valid}, {message}")
        
        # 清理临时文件
        os.unlink(temp_pdf_path)
        
        return is_valid
    except Exception as e:
        print(f"❌ 文件验证测试失败: {e}")
        return False

def test_text_to_pdf_conversion():
    """测试文本转PDF功能"""
    try:
        from apps.tools.pdf_converter_api import PDFConverter
        converter = PDFConverter()
        
        test_text = "这是一个测试文本。\nThis is a test text.\n包含中文和英文内容。"
        
        success, result, file_type = converter.text_to_pdf(test_text)
        
        if success:
            print(f"✅ 文本转PDF成功!")
            print(f"   输出类型: {file_type}")
            print(f"   输出大小: {len(result)} 字节")
            
            # 保存测试文件
            with open('test_text_to_pdf_output.pdf', 'wb') as f:
                f.write(result)
            print(f"   测试文件已保存: test_text_to_pdf_output.pdf")
            
            return True
        else:
            print(f"❌ 文本转PDF失败: {result}")
            return False
            
    except Exception as e:
        print(f"❌ 文本转PDF测试失败: {e}")
        return False

def test_pdf_to_text_conversion():
    """测试PDF转文本功能"""
    try:
        from apps.tools.pdf_converter_api import PDFConverter
        converter = PDFConverter()
        
        # 创建测试PDF
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_pdf_path = temp_file.name
        
        # 创建PDF内容
        c = canvas.Canvas(temp_pdf_path, pagesize=letter)
        c.drawString(100, 750, "Test PDF")
        c.drawString(100, 720, "This is a test PDF for conversion.")
        c.drawString(100, 690, "包含中文和英文内容。")
        c.save()
        
        # 模拟文件对象
        class MockFile:
            def __init__(self, path):
                self.name = os.path.basename(path)
                self.size = os.path.getsize(path)
                self.path = path
            
            def read(self):
                with open(self.path, 'rb') as f:
                    return f.read()
            
            def chunks(self):
                with open(self.path, 'rb') as f:
                    while True:
                        chunk = f.read(8192)
                        if not chunk:
                            break
                        yield chunk
        
        mock_file = MockFile(temp_pdf_path)
        
        success, result, file_type = converter.pdf_to_text(mock_file)
        
        if success:
            print(f"✅ PDF转文本成功!")
            print(f"   输出类型: {file_type}")
            print(f"   提取文本: {result[:100]}...")
            
            # 保存测试文件
            with open('test_pdf_to_text_output.txt', 'w', encoding='utf-8') as f:
                f.write(result)
            print(f"   测试文件已保存: test_pdf_to_text_output.txt")
            
            # 清理临时文件
            os.unlink(temp_pdf_path)
            
            return True
        else:
            print(f"❌ PDF转文本失败: {result}")
            # 清理临时文件
            os.unlink(temp_pdf_path)
            return False
            
    except Exception as e:
        print(f"❌ PDF转文本测试失败: {e}")
        return False

def test_image_to_pdf_conversion():
    """测试图片转PDF功能"""
    try:
        from apps.tools.pdf_converter_api import PDFConverter
        converter = PDFConverter()
        
        # 创建测试图片
        from PIL import Image, ImageDraw, ImageFont
        
        img = Image.new('RGB', (400, 300), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 20)
        except:
            font = ImageFont.load_default()
        
        draw.text((50, 50), "Test Image", fill='black', font=font)
        draw.text((50, 100), "测试图片", fill='black', font=font)
        
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
            temp_img_path = temp_file.name
        
        img.save(temp_img_path, 'PNG')
        
        # 模拟文件对象
        class MockFile:
            def __init__(self, path):
                self.name = os.path.basename(path)
                self.size = os.path.getsize(path)
                self.path = path
        
        mock_file = MockFile(temp_img_path)
        
        success, result, file_type = converter.images_to_pdf([mock_file])
        
        if success:
            print(f"✅ 图片转PDF成功!")
            print(f"   输出类型: {file_type}")
            print(f"   输出大小: {len(result)} 字节")
            
            # 保存测试文件
            with open('test_image_to_pdf_output.pdf', 'wb') as f:
                f.write(result)
            print(f"   测试文件已保存: test_image_to_pdf_output.pdf")
            
            # 清理临时文件
            os.unlink(temp_img_path)
            
            return True
        else:
            print(f"❌ 图片转PDF失败: {result}")
            # 清理临时文件
            os.unlink(temp_img_path)
            return False
            
    except Exception as e:
        print(f"❌ 图片转PDF测试失败: {e}")
        return False

def test_pdf_to_word_conversion():
    """测试PDF转Word功能"""
    try:
        from apps.tools.pdf_converter_api import PDFConverter
        converter = PDFConverter()
        
        # 创建测试PDF
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_pdf_path = temp_file.name
        
        # 创建PDF内容
        c = canvas.Canvas(temp_pdf_path, pagesize=letter)
        c.drawString(100, 750, "Test PDF for Word Conversion")
        c.drawString(100, 720, "This is a test PDF for Word conversion.")
        c.drawString(100, 690, "包含中文和英文内容。")
        c.save()
        
        # 模拟文件对象
        class MockFile:
            def __init__(self, path):
                self.name = os.path.basename(path)
                self.size = os.path.getsize(path)
                self.path = path
            
            def read(self):
                with open(self.path, 'rb') as f:
                    return f.read()
        
        mock_file = MockFile(temp_pdf_path)
        
        success, result, file_type = converter.pdf_to_word(mock_file)
        
        if success:
            print(f"✅ PDF转Word成功!")
            print(f"   输出类型: {file_type}")
            print(f"   输出大小: {len(result)} 字节")
            
            # 保存测试文件
            with open('test_pdf_to_word_output.docx', 'wb') as f:
                f.write(result)
            print(f"   测试文件已保存: test_pdf_to_word_output.docx")
            
            # 清理临时文件
            os.unlink(temp_pdf_path)
            
            return True
        else:
            print(f"❌ PDF转Word失败: {result}")
            # 清理临时文件
            os.unlink(temp_pdf_path)
            return False
            
    except Exception as e:
        print(f"❌ PDF转Word测试失败: {e}")
        return False

def test_word_to_pdf_conversion():
    """测试Word转PDF功能"""
    try:
        from apps.tools.pdf_converter_api import PDFConverter
        converter = PDFConverter()
        
        # 创建测试Word文档
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Word Document', 0)
        doc.add_paragraph('This is a test Word document for PDF conversion.')
        doc.add_paragraph('包含中文和英文内容。')
        doc.add_paragraph('Test Word to PDF conversion.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_docx_path = temp_file.name
        
        doc.save(temp_docx_path)
        
        # 模拟文件对象
        class MockFile:
            def __init__(self, path):
                self.name = os.path.basename(path)
                self.size = os.path.getsize(path)
                self.path = path
            
            def read(self):
                with open(self.path, 'rb') as f:
                    return f.read()
        
        mock_file = MockFile(temp_docx_path)
        
        success, result, file_type = converter.word_to_pdf(mock_file)
        
        if success:
            print(f"✅ Word转PDF成功!")
            print(f"   输出类型: {file_type}")
            print(f"   输出大小: {len(result)} 字节")
            
            # 保存测试文件
            with open('test_word_to_pdf_output.pdf', 'wb') as f:
                f.write(result)
            print(f"   测试文件已保存: test_word_to_pdf_output.pdf")
            
            # 清理临时文件
            os.unlink(temp_docx_path)
            
            return True
        else:
            print(f"❌ Word转PDF失败: {result}")
            # 清理临时文件
            os.unlink(temp_docx_path)
            return False
            
    except Exception as e:
        print(f"❌ Word转PDF测试失败: {e}")
        return False

def main():
    """主测试函数"""
    print("🚀 PDF转换引擎核心功能测试")
    print("=" * 60)
    
    tests = [
        ("PDF转换器导入", test_pdf_converter_import),
        ("PDF转换器初始化", test_pdf_converter_initialization),
        ("文件验证功能", test_file_validation),
        ("文本转PDF", test_text_to_pdf_conversion),
        ("PDF转文本", test_pdf_to_text_conversion),
        ("图片转PDF", test_image_to_pdf_conversion),
        ("PDF转Word", test_pdf_to_word_conversion),
        ("Word转PDF", test_word_to_pdf_conversion),
    ]
    
    results = {}
    
    for test_name, test_func in tests:
        print(f"\n🧪 测试 {test_name}...")
        try:
            result = test_func()
            results[test_name] = result
        except Exception as e:
            print(f"❌ {test_name} 测试异常: {e}")
            results[test_name] = False
    
    # 输出测试结果总结
    print("\n" + "=" * 60)
    print("📊 测试结果总结")
    print("=" * 60)
    
    total_tests = len(results)
    successful_tests = sum(results.values())
    
    for test_name, success in results.items():
        status = "✅ 成功" if success else "❌ 失败"
        print(f"{test_name:20} : {status}")
    
    print(f"\n总计: {successful_tests}/{total_tests} 个功能测试通过")
    
    if successful_tests == total_tests:
        print("🎉 所有PDF转换功能测试通过！")
        print("📁 测试文件已生成，请检查以下文件:")
        test_files = [
            'test_text_to_pdf_output.pdf',
            'test_pdf_to_text_output.txt',
            'test_image_to_pdf_output.pdf',
            'test_pdf_to_word_output.docx',
            'test_word_to_pdf_output.pdf'
        ]
        for file in test_files:
            if os.path.exists(file):
                print(f"   ✅ {file}")
            else:
                print(f"   ❌ {file} (未生成)")
    else:
        print("⚠️ 部分功能测试失败，请检查相关依赖和配置")
    
    return results

if __name__ == "__main__":
    main()
