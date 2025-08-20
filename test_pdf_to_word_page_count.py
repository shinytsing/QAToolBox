#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试PDF转Word页数问题修复
"""

import os
import sys
import django
import tempfile
import io

# 设置Django环境
sys.path.append('/Users/gaojie/PycharmProjects/QAToolBox')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')
django.setup()

from apps.tools.pdf_converter_api import PDFConverter
from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document

def create_test_pdf_with_pages(num_pages=2):
    """创建一个测试PDF文件，指定页数"""
    # 这是一个简化的PDF内容，实际使用时可能需要更复杂的PDF生成
    pdf_content = b"%PDF-1.4\n"
    
    # 添加页面内容
    for page in range(num_pages):
        page_obj = f"{page+1} 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/MediaBox [0 0 612 792]\n/Contents {page+3} 0 R\n>>\nendobj\n".encode('utf-8')
        content_obj = f"{page+3} 0 obj\n<<\n/Length 50\n>>\nstream\nBT\n/F1 12 Tf\n72 720 Td\n(Page {page+1} Content) Tj\nET\nendstream\nendobj\n".encode('utf-8')
        pdf_content += page_obj + content_obj
    
    # 添加页面树
    kids = " ".join([f"{i+1} 0 R" for i in range(num_pages)])
    pages_obj = f"2 0 obj\n<<\n/Type /Pages\n/Kids [{kids}]\n/Count {num_pages}\n>>\nendobj\n".encode('utf-8')
    pdf_content += pages_obj
    
    # 添加目录
    catalog_obj = "1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n".encode('utf-8')
    pdf_content += catalog_obj
    
    # 添加交叉引用表和尾部
    xref_offset = len(pdf_content)
    xref_header = "xref\n".encode('utf-8')
    xref_count = f"0 {num_pages*2+3}\n".encode('utf-8')
    xref_null = "0000000000 65535 f \n".encode('utf-8')
    pdf_content += xref_header + xref_count + xref_null
    
    # 添加对象偏移量
    offset = 0
    for i in range(num_pages*2+2):
        offset_str = f"{offset:010d} 00000 n \n".encode('utf-8')
        pdf_content += offset_str
        if i == 0:
            offset += 50  # 目录大小
        elif i == 1:
            offset += 50 + num_pages*20  # 页面树大小
        else:
            offset += 100  # 页面对象大小
    
    trailer_start = "trailer\n<<\n".encode('utf-8')
    size_str = f"/Size {num_pages*2+3}\n".encode('utf-8')
    root_str = "/Root 1 0 R\n".encode('utf-8')
    trailer_end = ">>\n".encode('utf-8')
    startxref_str = f"startxref\n{xref_offset}\n%%EOF".encode('utf-8')
    
    pdf_content += trailer_start + size_str + root_str + trailer_end + startxref_str
    
    return pdf_content

def test_pdf_to_word_page_count():
    """测试PDF转Word的页数保持"""
    print("🔍 测试PDF转Word页数保持...")
    converter = PDFConverter()
    
    # 创建2页的测试PDF
    test_pdf_content = create_test_pdf_with_pages(2)
    pdf_file = SimpleUploadedFile("test_2pages.pdf", test_pdf_content, content_type="application/pdf")
    
    try:
        print("📄 原始PDF: 2页")
        
        # 执行转换
        success, result, file_type = converter.pdf_to_word(pdf_file)
        
        if success:
            print("✅ PDF转Word成功")
            
            # 检查Word文档的页数
            doc = Document(io.BytesIO(result))
            
            # 计算页数（通过分页符数量+1）
            page_count = 1
            for paragraph in doc.paragraphs:
                if paragraph._element.xml.find('w:br') != -1 and 'w:type="page"' in paragraph._element.xml:
                    page_count += 1
            
            print(f"📄 转换后Word文档页数: {page_count}")
            
            if page_count == 2:
                print("✅ 页数保持正确！")
            else:
                print(f"⚠️ 页数发生变化: 2页 -> {page_count}页")
                
                # 分析内容结构
                print("\n📊 内容分析:")
                print(f"段落数量: {len(doc.paragraphs)}")
                
                # 检查是否有页面分隔符
                page_breaks = 0
                for paragraph in doc.paragraphs:
                    if paragraph._element.xml.find('w:br') != -1:
                        page_breaks += 1
                
                print(f"页面分隔符数量: {page_breaks}")
                
        else:
            print(f"❌ PDF转Word失败: {result}")
            
    except Exception as e:
        print(f"❌ 测试异常: {str(e)}")
    finally:
        pdf_file.close()

def test_ocr_page_structure():
    """测试OCR处理的页面结构"""
    print("\n🔍 测试OCR页面结构处理...")
    
    # 模拟OCR识别的页面文本
    ocr_texts = [
        "第一页内容\n这是第一页的第一行\n这是第一页的第二行",
        "第二页内容\n这是第二页的第一行\n这是第二页的第二行"
    ]
    
    print("📄 原始OCR文本:")
    for i, text in enumerate(ocr_texts):
        print(f"  页面{i+1}: {text}")
    
    # 模拟修复后的处理逻辑
    from docx import Document
    from docx.shared import Pt
    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 为每个页面创建单独的段落，保持页面结构
    for page_index, page_text in enumerate(ocr_texts):
        if page_text.strip():  # 只处理非空页面
            # 添加页面分隔符（除了第一页）
            if page_index > 0:
                document.add_page_break()
            
            # 将页面文本按行分割，保持原始格式
            lines = page_text.strip().split('\n')
            for line in lines:
                if line.strip():  # 只添加非空行
                    p = document.add_paragraph()
                    p.add_run(line.strip())
    
    print(f"📄 生成的Word文档段落数: {len(document.paragraphs)}")
    
    # 保存到临时文件进行测试
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        document.save(temp_file.name)
        temp_path = temp_file.name
    
    try:
        # 重新读取文档检查结构
        doc = Document(temp_path)
        print(f"📄 重新读取的Word文档段落数: {len(doc.paragraphs)}")
        
        # 显示段落内容
        print("📄 段落内容:")
        for i, para in enumerate(doc.paragraphs):
            print(f"  段落{i+1}: {para.text}")
            
    finally:
        # 清理临时文件
        try:
            os.unlink(temp_path)
        except:
            pass

if __name__ == "__main__":
    test_pdf_to_word_page_count()
    test_ocr_page_structure()
    print("\n✅ 测试完成")
