#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试PDF转换器的真实实现
"""

import os
import sys
import django
import tempfile
import io

# 设置Django环境
sys.path.append('/Users/gaojie/PycharmProjects/QAToolBox')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')
django.setup()

from apps.tools.pdf_converter_api import PDFConverter
from django.core.files.uploadedfile import SimpleUploadedFile

def create_real_test_pdf():
    """创建一个真实的测试PDF文件"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        
        # 创建临时PDF文件
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        # 使用reportlab创建PDF
        c = canvas.Canvas(temp_pdf_path, pagesize=A4)
        
        # 添加第一页内容
        c.drawString(100, 750, "第一页内容")
        c.drawString(100, 720, "这是第一页的第一行文本")
        c.drawString(100, 690, "这是第一页的第二行文本")
        c.drawString(100, 660, "这是第一页的第三行文本")
        
        # 添加第二页
        c.showPage()
        c.drawString(100, 750, "第二页内容")
        c.drawString(100, 720, "这是第二页的第一行文本")
        c.drawString(100, 690, "这是第二页的第二行文本")
        c.drawString(100, 660, "这是第二页的第三行文本")
        
        c.save()
        
        # 读取PDF内容
        with open(temp_pdf_path, 'rb') as f:
            pdf_content = f.read()
        
        # 清理临时文件
        os.unlink(temp_pdf_path)
        
        return pdf_content
        
    except ImportError:
        print("❌ reportlab未安装，无法创建测试PDF")
        return None
    except Exception as e:
        print(f"❌ 创建测试PDF失败: {e}")
        return None

def test_pdf_converter_real():
    """测试PDF转换器的真实实现"""
    print("🔍 测试PDF转换器真实实现...")
    
    # 创建真实的测试PDF
    pdf_content = create_real_test_pdf()
    if not pdf_content:
        print("❌ 无法创建测试PDF，跳过测试")
        return
    
    print(f"✅ 创建测试PDF成功，大小: {len(pdf_content)} bytes")
    
    # 创建PDF文件对象
    pdf_file = SimpleUploadedFile("test_real.pdf", pdf_content, content_type="application/pdf")
    
    # 创建转换器实例
    converter = PDFConverter()
    
    try:
        print("📄 开始PDF转Word转换...")
        
        # 执行转换
        success, result, file_type = converter.pdf_to_word(pdf_file)
        
        if success:
            print("✅ PDF转Word转换成功!")
            print(f"   文件类型: {file_type}")
            print(f"   结果大小: {len(result)} bytes")
            
            # 检查结果是否为真实的Word文档
            if result.startswith(b'PK'):  # ZIP文件头，DOCX是ZIP格式
                print("✅ 结果确实是Word文档格式")
                
                # 尝试解析Word文档内容
                try:
                    from docx import Document
                    doc = Document(io.BytesIO(result))
                    
                    print(f"   段落数量: {len(doc.paragraphs)}")
                    print("   文档内容预览:")
                    for i, para in enumerate(doc.paragraphs[:5]):  # 显示前5段
                        if para.text.strip():
                            print(f"     段落{i+1}: {para.text[:50]}...")
                    
                except Exception as e:
                    print(f"⚠️ 无法解析Word文档内容: {e}")
            else:
                print("❌ 结果不是Word文档格式")
                print(f"   文件头: {result[:20]}")
        else:
            print(f"❌ PDF转Word转换失败: {result}")
            
    except Exception as e:
        print(f"❌ 转换过程异常: {str(e)}")
    finally:
        pdf_file.close()

def test_text_to_pdf_real():
    """测试文本转PDF的真实实现"""
    print("\n🔍 测试文本转PDF真实实现...")
    
    converter = PDFConverter()
    test_text = "这是一个测试文本。\n包含多行内容。\n用于测试文本转PDF功能。"
    
    try:
        print("📄 开始文本转PDF转换...")
        
        # 执行转换
        success, result, file_type = converter.text_to_pdf(test_text)
        
        if success:
            print("✅ 文本转PDF转换成功!")
            print(f"   文件类型: {file_type}")
            print(f"   结果大小: {len(result)} bytes")
            
            # 检查结果是否为真实的PDF
            if result.startswith(b'%PDF'):  # PDF文件头
                print("✅ 结果确实是PDF格式")
            else:
                print("❌ 结果不是PDF格式")
                print(f"   文件头: {result[:20]}")
        else:
            print(f"❌ 文本转PDF转换失败: {result}")
            
    except Exception as e:
        print(f"❌ 转换过程异常: {str(e)}")

def test_converter_methods():
    """测试转换器的方法"""
    print("\n🔍 测试转换器方法...")
    
    converter = PDFConverter()
    
    # 检查支持的方法
    methods = [method for method in dir(converter) if not method.startswith('_')]
    print(f"转换器方法: {methods}")
    
    # 检查支持的格式
    print(f"支持的格式: {converter.supported_formats}")

if __name__ == "__main__":
    test_converter_methods()
    test_pdf_converter_real()
    test_text_to_pdf_real()
    print("\n✅ 测试完成")
