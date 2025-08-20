#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试Web API的Word转PDF中文字符处理
"""

import os
import sys
import django
import json
import base64

# 设置Django环境
sys.path.append('/Users/gaojie/PycharmProjects/QAToolBox')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')
django.setup()

from django.test import Client
from django.contrib.auth.models import User
from docx import Document
import io

def test_web_api_word_to_pdf():
    """测试Web API的Word转PDF功能"""
    print("🔍 测试Web API Word转PDF中文字符处理...")
    
    try:
        # 创建测试用户
        client = Client()
        user, created = User.objects.get_or_create(
            username='test_user',
            defaults={'email': 'test@example.com'}
        )
        if created:
            user.set_password('test_password')
            user.save()
        
        # 登录
        client.force_login(user)
        
        # 创建包含中文的Word文档
        doc = Document()
        doc.add_heading('中文测试文档', 0)
        doc.add_paragraph('这是一个包含中文内容的测试文档。')
        doc.add_paragraph('测试内容包括：')
        doc.add_paragraph('1. 中文字符显示')
        doc.add_paragraph('2. 特殊符号：！@#￥%……&*（）')
        doc.add_paragraph('3. 数字：1234567890')
        doc.add_paragraph('4. 英文：Hello World')
        doc.add_paragraph('5. 混合内容：中文123ABC！@#')
        doc.add_paragraph('')
        doc.add_paragraph('这是一个较长的段落，用来测试中文字符在PDF中的显示效果。包含各种标点符号和特殊字符，确保转换后的PDF能够正确显示所有内容。')
        doc.add_paragraph('测试完成！')
        
        # 保存到内存
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_content = doc_buffer.getvalue()
        doc_buffer.close()
        
        print("   正在调用Web API...")
        
        # 构造文件上传数据
        from django.core.files.uploadedfile import SimpleUploadedFile
        uploaded_file = SimpleUploadedFile(
            "chinese_test.docx",
            doc_content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # 调用API (使用Django测试客户端的正确格式)
        response = client.post(
            '/tools/api/pdf-converter/',
            data={
                'type': 'word-to-pdf',
                'file': uploaded_file
            }
        )
        
        print(f"   API响应状态码: {response.status_code}")
        
        if response.status_code == 200:
            response_data = response.json()
            if response_data.get('success'):
                print("✅ Web API Word转PDF成功")
                print(f"   文件名: {response_data.get('filename', 'N/A')}")
                print(f"   下载URL: {response_data.get('download_url', 'N/A')}")
            else:
                print(f"❌ Web API Word转PDF失败: {response_data.get('error', 'Unknown error')}")
        else:
            print(f"❌ Web API请求失败: HTTP {response.status_code}")
            try:
                error_data = response.json()
                print(f"   错误信息: {error_data.get('error', 'Unknown error')}")
            except:
                print(f"   响应内容: {response.content.decode('utf-8')}")
        
    except Exception as e:
        print(f"❌ 测试异常: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("🚀 开始测试Web API Word转PDF...")
    print("=" * 50)
    
    test_web_api_word_to_pdf()
    
    print("=" * 50)
    print("🎉 测试完成！")
