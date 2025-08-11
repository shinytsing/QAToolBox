#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试Word和图片转换功能
"""

import os
import sys
import django

# 设置Django环境
sys.path.append('/Users/gaojie/PycharmProjects/QAToolBox')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')
django.setup()

from apps.tools.pdf_converter_api import PDFConverter
from django.core.files.uploadedfile import SimpleUploadedFile
import io

def test_pdf_to_word():
    """测试PDF转Word"""
    print("🔍 测试PDF转Word...")
    converter = PDFConverter()
    
    # 创建测试PDF内容
    test_content = b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/MediaBox [0 0 612 792]\n/Contents 4 0 R\n>>\nendobj\n4 0 obj\n<<\n/Length 44\n>>\nstream\nBT\n/F1 12 Tf\n72 720 Td\n(Hello World) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000204 00000 n \ntrailer\n<<\n/Size 5\n/Root 1 0 R\n>>\nstartxref\n297\n%%EOF"
    
    pdf_file = SimpleUploadedFile("test.pdf", test_content, content_type="application/pdf")
    
    try:
        success, result, file_type = converter.pdf_to_word(pdf_file)
        if success:
            print("✅ PDF转Word成功")
            print(f"   生成Word文件大小: {len(result)} bytes")
        else:
            print(f"❌ PDF转Word失败: {result}")
    except Exception as e:
        print(f"❌ PDF转Word异常: {str(e)}")
    finally:
        pdf_file.close()

def test_word_to_pdf():
    """测试Word转PDF"""
    print("🔍 测试Word转PDF...")
    converter = PDFConverter()
    
    try:
        # 创建简单Word文档
        from docx import Document
        doc = Document()
        doc.add_paragraph("这是一个测试Word文档")
        doc.add_paragraph("包含多行内容")
        
        # 保存到内存
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        word_content = doc_buffer.getvalue()
        word_file = SimpleUploadedFile("test.docx", word_content, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        success, result, file_type = converter.word_to_pdf(word_file)
        if success:
            print("✅ Word转PDF成功")
            print(f"   生成PDF大小: {len(result)} bytes")
        else:
            print(f"❌ Word转PDF失败: {result}")
            
        word_file.close()
        doc_buffer.close()
        
    except ImportError:
        print("⚠️  python-docx库未安装，跳过Word转PDF测试")
    except Exception as e:
        print(f"❌ Word转PDF异常: {str(e)}")

def test_pdf_to_images():
    """测试PDF转图片"""
    print("🔍 测试PDF转图片...")
    converter = PDFConverter()
    
    # 创建测试PDF内容
    test_content = b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/MediaBox [0 0 612 792]\n/Contents 4 0 R\n>>\nendobj\n4 0 obj\n<<\n/Length 44\n>>\nstream\nBT\n/F1 12 Tf\n72 720 Td\n(Hello World) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000204 00000 n \ntrailer\n<<\n/Size 5\n/Root 1 0 R\n>>\nstartxref\n297\n%%EOF"
    
    pdf_file = SimpleUploadedFile("test.pdf", test_content, content_type="application/pdf")
    
    try:
        success, result, file_type = converter.pdf_to_images(pdf_file)
        if success:
            print("✅ PDF转图片成功")
            print(f"   生成图片数量: {len(result)}")
        else:
            print(f"❌ PDF转图片失败: {result}")
    except Exception as e:
        print(f"❌ PDF转图片异常: {str(e)}")
    finally:
        pdf_file.close()

def test_images_to_pdf():
    """测试图片转PDF"""
    print("🔍 测试图片转PDF...")
    converter = PDFConverter()
    
    try:
        # 创建简单图片
        from PIL import Image, ImageDraw
        img = Image.new('RGB', (200, 100), color='white')
        draw = ImageDraw.Draw(img)
        draw.text((10, 40), "Test Image", fill='black')
        
        # 保存到内存
        img_buffer = io.BytesIO()
        img.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        img_content = img_buffer.getvalue()
        img_file = SimpleUploadedFile("test.png", img_content, content_type="image/png")
        
        success, result, file_type = converter.images_to_pdf([img_file])
        if success:
            print("✅ 图片转PDF成功")
            print(f"   生成PDF大小: {len(result)} bytes")
        else:
            print(f"❌ 图片转PDF失败: {result}")
            
        img_file.close()
        img_buffer.close()
        
    except ImportError:
        print("⚠️  Pillow库未安装，跳过图片转PDF测试")
    except Exception as e:
        print(f"❌ 图片转PDF异常: {str(e)}")

def test_batch_conversion():
    """测试批量转换"""
    print("🔍 测试批量转换...")
    converter = PDFConverter()
    
    # 创建多个测试文件
    files = []
    
    # PDF文件
    pdf_content = b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/MediaBox [0 0 612 792]\n/Contents 4 0 R\n>>\nendobj\n4 0 obj\n<<\n/Length 44\n>>\nstream\nBT\n/F1 12 Tf\n72 720 Td\n(Hello World) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000204 00000 n \ntrailer\n<<\n/Size 5\n/Root 1 0 R\n>>\nstartxref\n297\n%%EOF"
    pdf_file = SimpleUploadedFile("test1.pdf", pdf_content, content_type="application/pdf")
    files.append(pdf_file)
    
    # TXT文件
    txt_content = "这是第二个测试文件\n包含文本内容".encode('utf-8')
    txt_file = SimpleUploadedFile("test2.txt", txt_content, content_type="text/plain")
    files.append(txt_file)
    
    try:
        print("   测试批量PDF转文本...")
        for i, file in enumerate(files):
            if file.name.endswith('.pdf'):
                success, result, file_type = converter.pdf_to_text(file)
                if success:
                    print(f"   ✅ 文件{i+1}转换成功")
                else:
                    print(f"   ❌ 文件{i+1}转换失败: {result}")
                    
    except Exception as e:
        print(f"❌ 批量转换异常: {str(e)}")
    finally:
        for file in files:
            file.close()

def main():
    """主测试函数"""
    print("🚀 开始测试Word和图片转换功能...")
    print("=" * 50)
    
    # 测试所有转换功能
    test_pdf_to_word()
    print()
    
    test_word_to_pdf()
    print()
    
    test_pdf_to_images()
    print()
    
    test_images_to_pdf()
    print()
    
    test_batch_conversion()
    print()
    
    print("=" * 50)
    print("🎉 所有测试完成！")

if __name__ == "__main__":
    main() 