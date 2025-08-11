#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word转PDF临时文件问题测试和修复脚本
"""

import os
import sys
import tempfile
import io
import shutil
from pathlib import Path

# 设置Django环境
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')

def test_word_to_pdf_issue():
    """测试Word转PDF的临时文件问题"""
    print("🧪 测试Word转PDF临时文件问题...")
    
    try:
        import django
        django.setup()
        
        from apps.tools.pdf_converter_api import PDFConverter
        from docx import Document
        
        # 创建一个简单的测试Word文档
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个测试Word文档。')
        doc.add_paragraph('包含中文和English内容。')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        try:
            # 创建文件对象
            with open(temp_docx_path, 'rb') as f:
                word_file = io.BytesIO(f.read())
                word_file.name = "test.docx"
            
            # 测试转换
            converter = PDFConverter()
            success, result, file_type = converter.word_to_pdf(word_file)
            
            if success:
                print("✅ Word转PDF转换成功!")
                print(f"   输出类型: {file_type}")
                print(f"   输出大小: {len(result)} 字节")
                return True
            else:
                print(f"❌ Word转PDF转换失败: {result}")
                return False
                
        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
                
    except Exception as e:
        print(f"❌ Word转PDF测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_docx2pdf_direct():
    """直接测试docx2pdf库"""
    print("\n🧪 直接测试docx2pdf库...")
    
    try:
        from docx2pdf import convert
        from docx import Document
        
        # 创建一个简单的测试Word文档
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个测试Word文档。')
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        try:
            print(f"   临时docx文件: {temp_docx_path}")
            print(f"   临时pdf文件: {temp_pdf_path}")
            
            # 测试转换
            convert(temp_docx_path, temp_pdf_path)
            
            # 检查输出文件
            if os.path.exists(temp_pdf_path) and os.path.getsize(temp_pdf_path) > 0:
                print("✅ docx2pdf转换测试成功")
                print(f"   输出文件大小: {os.path.getsize(temp_pdf_path)} 字节")
                return True
            else:
                print("❌ docx2pdf转换测试失败")
                return False
                
        except Exception as e:
            print(f"❌ docx2pdf转换测试失败: {e}")
            import traceback
            traceback.print_exc()
            return False
        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
                
    except ImportError as e:
        print(f"❌ docx2pdf导入失败: {e}")
        return False

def fix_word_to_pdf_method():
    """修复Word转PDF方法"""
    print("\n🔧 修复Word转PDF方法...")
    
    # 读取当前的PDF转换器文件
    pdf_converter_path = "apps/tools/pdf_converter_api.py"
    
    with open(pdf_converter_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 查找word_to_pdf方法
    if 'def word_to_pdf(self, word_file):' in content:
        print("✅ 找到word_to_pdf方法")
        
        # 创建修复后的方法
        fixed_method = '''    def word_to_pdf(self, word_file):
        """Word转PDF - 真实实现（修复版）"""
        try:
            if not DOCX2PDF_AVAILABLE:
                # 提供更详细的错误信息和解决方案
                error_msg = "docx2pdf库未安装，无法进行Word转PDF转换\\n"
                error_msg += "解决方案：\\n"
                error_msg += "1. 请确保已安装docx2pdf库：pip install docx2pdf\\n"
                error_msg += "2. 如果已安装，请重启服务器\\n"
                error_msg += "3. 检查Python环境是否正确"
                return False, error_msg, None
            
            # 重置文件指针
            word_file.seek(0)
            
            # 使用docx2pdf进行真实转换
            from io import BytesIO
            import tempfile
            import os
            import shutil
            
            # 创建临时目录
            temp_dir = tempfile.mkdtemp()
            
            try:
                # 创建临时输入文件
                temp_docx_path = os.path.join(temp_dir, 'input.docx')
                with open(temp_docx_path, 'wb') as temp_docx:
                    temp_docx.write(word_file.read())
                
                # 创建临时输出文件路径
                temp_pdf_path = os.path.join(temp_dir, 'output.pdf')
                
                # 使用docx2pdf进行转换
                from docx2pdf import convert
                convert(temp_docx_path, temp_pdf_path)
                
                # 检查输出文件是否存在
                if not os.path.exists(temp_pdf_path):
                    return False, "转换失败：输出PDF文件未生成", None
                
                # 读取转换后的文件
                with open(temp_pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                if len(pdf_content) == 0:
                    return False, "转换后的文件为空，可能是Word文档内容无法识别", None
                
                return True, pdf_content, "word_to_pdf"
                
            finally:
                # 清理临时目录
                try:
                    shutil.rmtree(temp_dir)
                except Exception as cleanup_error:
                    logger.warning(f"清理临时目录失败: {cleanup_error}")
            
        except Exception as e:
            logger.error(f"Word转PDF失败: {str(e)}")
            return False, f"转换失败: {str(e)}", None'''
        
        # 替换方法
        import re
        pattern = r'def word_to_pdf\(self, word_file\):.*?except Exception as e:.*?return False, f"转换失败: {str\(e\)}", None'
        replacement = fixed_method
        
        # 使用更精确的替换
        lines = content.split('\n')
        new_lines = []
        in_method = False
        method_start = 0
        brace_count = 0
        
        for i, line in enumerate(lines):
            if 'def word_to_pdf(self, word_file):' in line:
                in_method = True
                method_start = i
                new_lines.extend(lines[:i])
                new_lines.extend(fixed_method.split('\n'))
                continue
            
            if in_method:
                if '{' in line:
                    brace_count += line.count('{')
                if '}' in line:
                    brace_count -= line.count('}')
                
                if brace_count == 0 and line.strip() == '':
                    in_method = False
                    continue
            else:
                if i > method_start:
                    new_lines.append(line)
        
        # 如果替换失败，使用简单的方法
        if len(new_lines) == 0:
            print("⚠️  使用简单替换方法")
            # 简单的字符串替换
            old_method_start = content.find('def word_to_pdf(self, word_file):')
            if old_method_start != -1:
                # 找到方法结束位置
                brace_count = 0
                method_end = old_method_start
                for i in range(old_method_start, len(content)):
                    if content[i] == '{':
                        brace_count += 1
                    elif content[i] == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            method_end = i + 1
                            break
                
                new_content = content[:old_method_start] + fixed_method + content[method_end:]
                
                # 写回文件
                with open(pdf_converter_path, 'w', encoding='utf-8') as f:
                    f.write(new_content)
                
                print("✅ Word转PDF方法已修复")
                return True
        
        # 写回文件
        with open(pdf_converter_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(new_lines))
        
        print("✅ Word转PDF方法已修复")
        return True
    else:
        print("❌ 未找到word_to_pdf方法")
        return False

def main():
    """主函数"""
    print("🔍 Word转PDF临时文件问题诊断和修复")
    print("=" * 50)
    
    # 测试当前问题
    test_result = test_word_to_pdf_issue()
    
    # 直接测试docx2pdf
    direct_test_result = test_docx2pdf_direct()
    
    # 如果测试失败，尝试修复
    if not test_result or not direct_test_result:
        print("\n🔧 尝试修复Word转PDF方法...")
        fix_success = fix_word_to_pdf_method()
        
        if fix_success:
            print("\n🔄 重新测试修复后的方法...")
            test_result = test_word_to_pdf_issue()
    
    print("\n" + "=" * 50)
    if test_result:
        print("✅ Word转PDF问题已解决")
    else:
        print("❌ Word转PDF问题仍然存在")
        print("建议检查docx2pdf库的安装和配置")

if __name__ == "__main__":
    main() 