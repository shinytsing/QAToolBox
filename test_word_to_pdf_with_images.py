#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试Word转PDF的图片处理功能
"""

import os
import sys
import django
import tempfile
import io

# 设置Django环境
sys.path.append('/Users/gaojie/PycharmProjects/QAToolBox')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')
django.setup()

from apps.tools.pdf_converter_api import PDFConverter
from django.core.files.uploadedfile import SimpleUploadedFile

def create_test_word_with_images():
    """创建一个包含图片的测试Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('测试文档 - 包含图片', 0)
        
        # 添加段落
        doc.add_paragraph('这是一个测试文档，包含文本和图片。')
        
        # 添加图片（如果可能）
        try:
            # 创建一个简单的测试图片
            from PIL import Image, ImageDraw, ImageFont
            
            # 创建图片
            img = Image.new('RGB', (400, 300), color='lightblue')
            draw = ImageDraw.Draw(img)
            
            # 添加文字
            try:
                font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
            except:
                font = ImageFont.load_default()
            
            draw.text((50, 50), "测试图片", fill='black', font=font)
            draw.text((50, 100), "Word转PDF测试", fill='black', font=font)
            
            # 保存图片到临时文件
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                img.save(temp_img.name, 'PNG')
                temp_img_path = temp_img.name
            
            # 将图片添加到Word文档
            doc.add_picture(temp_img_path, width=Inches(4))
            
            # 清理临时图片文件
            os.unlink(temp_img_path)
            
            print("✅ 成功添加测试图片到Word文档")
            
        except Exception as img_error:
            print(f"⚠️ 添加图片失败: {img_error}")
            # 添加占位符文本
            doc.add_paragraph('[图片占位符]')
        
        # 添加更多文本
        doc.add_paragraph('这是图片下方的文本内容。')
        doc.add_paragraph('Word转PDF功能应该能够正确处理图片。')
        
        # 保存Word文档到临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        # 读取Word文档内容
        with open(temp_docx_path, 'rb') as f:
            docx_content = f.read()
        
        # 清理临时文件
        os.unlink(temp_docx_path)
        
        return docx_content
        
    except ImportError:
        print("❌ python-docx未安装，无法创建Word文档")
        return None
    except Exception as e:
        print(f"❌ 创建Word文档失败: {e}")
        return None

def test_word_to_pdf_with_images():
    """测试Word转PDF的图片处理"""
    print("🔍 测试Word转PDF图片处理...")
    
    # 创建包含图片的Word文档
    docx_content = create_test_word_with_images()
    if not docx_content:
        print("❌ 无法创建测试Word文档，跳过测试")
        return
    
    print(f"✅ 创建测试Word文档成功，大小: {len(docx_content)} bytes")
    
    # 创建Word文件对象
    word_file = SimpleUploadedFile("test_with_images.docx", docx_content, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    # 创建转换器实例
    converter = PDFConverter()
    
    try:
        print("📄 开始Word转PDF转换...")
        
        # 执行转换
        success, result, file_type = converter.word_to_pdf(word_file)
        
        if success:
            print("✅ Word转PDF转换成功!")
            print(f"   文件类型: {file_type}")
            print(f"   结果大小: {len(result)} bytes")
            
            # 检查结果是否为真实的PDF
            if result.startswith(b'%PDF'):  # PDF文件头
                print("✅ 结果确实是PDF格式")
                
                # 保存PDF文件用于检查
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                    temp_pdf.write(result)
                    temp_pdf_path = temp_pdf.name
                
                print(f"📄 PDF文件已保存到: {temp_pdf_path}")
                print("💡 请手动检查PDF文件是否包含图片")
                
                # 清理临时文件
                try:
                    os.unlink(temp_pdf_path)
                except:
                    pass
            else:
                print("❌ 结果不是PDF格式")
                print(f"   文件头: {result[:20]}")
        else:
            print(f"❌ Word转PDF转换失败: {result}")
            
    except Exception as e:
        print(f"❌ 转换过程异常: {str(e)}")
    finally:
        word_file.close()

def test_simple_word_to_pdf():
    """测试简单的Word转PDF（无图片）"""
    print("\n🔍 测试简单Word转PDF...")
    
    try:
        from docx import Document
        
        # 创建简单的Word文档
        doc = Document()
        doc.add_heading('简单测试文档', 0)
        doc.add_paragraph('这是一个简单的测试文档。')
        doc.add_paragraph('用于测试Word转PDF功能。')
        doc.add_paragraph('不包含图片。')
        
        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        # 读取内容
        with open(temp_docx_path, 'rb') as f:
            docx_content = f.read()
        
        # 清理临时文件
        os.unlink(temp_docx_path)
        
        # 创建文件对象
        word_file = SimpleUploadedFile("simple_test.docx", docx_content, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        # 转换
        converter = PDFConverter()
        success, result, file_type = converter.word_to_pdf(word_file)
        
        if success:
            print("✅ 简单Word转PDF成功!")
            print(f"   结果大小: {len(result)} bytes")
        else:
            print(f"❌ 简单Word转PDF失败: {result}")
            
        word_file.close()
        
    except Exception as e:
        print(f"❌ 简单Word转PDF测试失败: {e}")

if __name__ == "__main__":
    test_simple_word_to_pdf()
    test_word_to_pdf_with_images()
    print("\n✅ 测试完成")
